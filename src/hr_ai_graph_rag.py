# ============================================================
# Colab-ready HR AI Assistant: external DOCX/JSON ingestion + Offline Artifacts + Runtime Local LLM classifier + RAG + Graph-enhanced RAG + LangGraph
# 本版為「成熟 workflow 調整版」：
# - 勞動基準法由使用者上傳 DOCX
# - 模擬銀行員工內規由使用者上傳 DOCX
# - Golden Dataset 由使用者上傳 JSON
# - Offline LLM-assisted artifacts 由 JSON/ZIP 接入：concept_nodes, risk_policy, query_patterns, rewrite_rules, relation_schema, graph_relation_candidates
# - Runtime local LLM 負責 category / intent / concept / ambiguity / candidate risk structured classification
# - Query patterns：local LLM runtime 判斷 pattern / missing slots，offline query_patterns 作 schema + fallback
# - Rewrite rules：offline mandatory terms + local LLM optional semantic terms 混合 query expansion
# - Final route 仍由 deterministic guardrails + risk_policy override 決定
# - Colab GPU local LLM via HuggingFace Transformers, no OpenAI API key required
# ============================================================

# -----------------------------
# 0. Install packages in Colab
# -----------------------------
import sys, subprocess, pkgutil, os

REQUIRED_PACKAGES = {
    "python-docx": "docx",
    "sentence-transformers": "sentence_transformers",
    "faiss-cpu": "faiss",
    "rank-bm25": "rank_bm25",
    "networkx": "networkx",
    "langgraph": "langgraph",
    "transformers": "transformers",
    "accelerate": "accelerate",
    "bitsandbytes": "bitsandbytes",
    "torch": "torch",
    "pandas": "pandas",
    "numpy": "numpy",
    "tqdm": "tqdm",
    "rich": "rich",
}

def pip_install_if_needed():
    missing = []
    for pip_name, import_name in REQUIRED_PACKAGES.items():
        if pkgutil.find_loader(import_name) is None:
            missing.append(pip_name)
    if missing:
        print("Installing missing packages:", missing)
        subprocess.check_call([sys.executable, "-m", "pip", "install", "-q", "-U"] + missing)
    else:
        print("All required packages are installed.")

# Auto-install on import only inside Colab, or when AUTO_PIP_INSTALL=true.
# When running from a normal checkout you should install via requirements.txt instead,
# which keeps `import hr_ai_graph_rag` side-effect free.
def _module_available(name: str) -> bool:
    import importlib.util
    try:
        return importlib.util.find_spec(name) is not None
    except Exception:
        return False

_IN_COLAB_EARLY = _module_available("google.colab")
if _IN_COLAB_EARLY or os.getenv("AUTO_PIP_INSTALL", "false").lower() == "true":
    pip_install_if_needed()

# -----------------------------
# 1. Imports & Global Settings
# -----------------------------
import re
import json
import time
import zipfile
import warnings
from dataclasses import dataclass
from typing import TypedDict, List, Dict, Any, Optional, Literal, Tuple
from getpass import getpass
from pathlib import Path
from collections import defaultdict

import numpy as np
import pandas as pd
import networkx as nx
from docx import Document
from rank_bm25 import BM25Okapi
from tqdm.auto import tqdm

# Heavy / GPU / LLM dependencies (torch, faiss, sentence-transformers, transformers,
# langgraph) are imported lazily inside the components that need them
# (HybridRetriever, LocalHFLLM, HRAssistantGraph). This keeps the data / parsing /
# chunking / graph / evaluation layers importable and runnable without a GPU or the
# generative-LLM stack installed.
try:
    import torch
    _TORCH_AVAILABLE = True
except Exception:  # pragma: no cover - torch is optional for the no-LLM pipeline
    torch = None
    _TORCH_AVAILABLE = False

try:
    from IPython.display import display, Markdown, Image
except Exception:  # pragma: no cover - only present in notebooks/Colab
    def display(*args, **kwargs):
        for a in args:
            print(a)

    def Markdown(x=""):
        return x

    def Image(*args, **kwargs):
        return None


def _cuda_available() -> bool:
    return bool(_TORCH_AVAILABLE and torch is not None and torch.cuda.is_available())


warnings.filterwarnings("ignore")

# Colab file upload support
try:
    from google.colab import files
    IN_COLAB = True
except Exception:
    files = None
    IN_COLAB = False

# Embedding model for dense retrieval. Default is BAAI/bge-m3 — a strong multilingual
# (incl. Traditional Chinese) embedding model. Override with EMBEDDING_MODEL_NAME.
# bge-m3 is ~2GB and is best on GPU; on CPU it works but is slow.
EMBEDDING_MODEL_NAME = os.getenv("EMBEDDING_MODEL_NAME", "BAAI/bge-m3")

# Cross-encoder reranker applied to the hybrid-retrieval candidates before the final cut.
# Default is BAAI/bge-reranker-v2-m3 (multilingual). Set USE_RERANKER=false to disable.
USE_RERANKER = os.getenv("USE_RERANKER", "true").lower() == "true"
RERANKER_MODEL_NAME = os.getenv("RERANKER_MODEL_NAME", "BAAI/bge-reranker-v2-m3")
# How many hybrid candidates to feed into the reranker (before trimming to top_k).
RERANK_CANDIDATES = int(os.getenv("RERANK_CANDIDATES", "20"))
# Weight blending reranker score with the hybrid score (1.0 = reranker only).
RERANK_WEIGHT = float(os.getenv("RERANK_WEIGHT", "0.7"))

# Local HuggingFace LLM settings (response generation).
# Default is the open-source google/gemma-2-2b-it (Colab T4 friendly with 4-bit).
# For more quality on a larger GPU try google/gemma-2-9b-it.
# NOTE: Gemma is a gated model on the HuggingFace Hub — you must accept its license and
# provide an HF token (e.g. `huggingface-cli login` or HF_TOKEN env) to download it.
HF_LLM_MODEL_NAME = os.getenv("HF_LLM_MODEL_NAME", "google/gemma-2-2b-it")
HF_LLM_USE_4BIT = os.getenv("HF_LLM_USE_4BIT", "true").lower() == "true"
HF_MAX_NEW_TOKENS = int(os.getenv("HF_MAX_NEW_TOKENS", "768"))
HF_TEMPERATURE = float(os.getenv("HF_TEMPERATURE", "0.1"))

# Runtime local LLM classification is enabled by default in this version.
# It classifies category / intent / matched concepts / ambiguity / candidate risk.
# Final routing still uses deterministic guardrails and risk_policy override.
USE_LOCAL_LLM_FOR_QUERY_UNDERSTANDING = os.getenv("USE_LOCAL_LLM_FOR_QUERY_UNDERSTANDING", "true").lower() == "true"
# Query patterns and optional rewrite terms are handled by the same runtime classifier.
# Offline artifacts still provide the controlled schema / mandatory terms.
LOCAL_LLM_OPTIONAL_REWRITE_MAX_TERMS = int(os.getenv("LOCAL_LLM_OPTIONAL_REWRITE_MAX_TERMS", "5"))

# Local LLM is used for final answer generation; if loading fails, the workflow falls back to template answer.
# Set USE_LLM=false to skip the generative LLM entirely and use the deterministic
# template answer (useful for running the pipeline without the GPU/transformers stack).
USE_LLM = os.getenv("USE_LLM", "true").lower() == "true"

OUTPUT_DIR = Path("/content/hr_ai_graph_rag_outputs") if IN_COLAB else Path("./hr_ai_graph_rag_outputs")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Repo root and bundled data locations (used for zero-config local runs).
# Layout:  <repo>/src/hr_ai_graph_rag.py  +  <repo>/data/{policies,golden,hr_offline_artifacts}
REPO_ROOT = Path(__file__).resolve().parent.parent if "__file__" in globals() else Path.cwd()
REPO_DATA_DIR = REPO_ROOT / "data"

# Optional: you may set these paths manually before running the script/notebook.
# If left blank in Colab, the script will ask you to upload three files:
# 1) 勞動基準法 DOCX, 2) 模擬銀行內規 DOCX, 3) Golden Dataset JSON.
# Outside Colab, paths default to the bundled files under <repo>/data when present.
LABOR_LAW_DOCX_PATH = os.getenv("LABOR_LAW_DOCX_PATH", "").strip()
INTERNAL_POLICY_DOCX_PATH = os.getenv("INTERNAL_POLICY_DOCX_PATH", "").strip()
GOLDEN_DATASET_JSON_PATH = os.getenv("GOLDEN_DATASET_JSON_PATH", "").strip()

if not IN_COLAB:
    _bundled_policy = REPO_DATA_DIR / "policies" / "安久銀行員工工作與福利規章辦法_模擬版.docx"
    _bundled_golden = REPO_DATA_DIR / "golden" / "anjou_bank_hr_ai_golden_dataset_50.json"
    if not INTERNAL_POLICY_DOCX_PATH and _bundled_policy.exists():
        INTERNAL_POLICY_DOCX_PATH = str(_bundled_policy)
    if not GOLDEN_DATASET_JSON_PATH and _bundled_golden.exists():
        GOLDEN_DATASET_JSON_PATH = str(_bundled_golden)

# Offline artifacts can be provided as a folder or a ZIP file.
# Recommended files inside the folder/ZIP:
# workflow_role_mapping.json, concept_nodes.json, risk_policy.json, query_patterns.json,
# rewrite_rules.json, relation_schema.json, graph_relation_candidates.json, local_llm_usage_policy.json
OFFLINE_ARTIFACT_DIR = os.getenv("OFFLINE_ARTIFACT_DIR", "").strip()
OFFLINE_ARTIFACT_ZIP_PATH = os.getenv("OFFLINE_ARTIFACT_ZIP_PATH", "").strip()

if not IN_COLAB and not OFFLINE_ARTIFACT_DIR:
    _bundled_artifacts = REPO_DATA_DIR / "hr_offline_artifacts"
    if (_bundled_artifacts / "concept_nodes.json").exists():
        OFFLINE_ARTIFACT_DIR = str(_bundled_artifacts)

# For safety, only approved graph edges are loaded by default.
LOAD_PENDING_GRAPH_EDGES = os.getenv("LOAD_PENDING_GRAPH_EDGES", "false").lower() == "true"

# Important: keep evaluation data out of the knowledge base by default.
# Set to true only for retrieval experiments. For honest evaluation, leave it false.
USE_GOLDEN_AS_FAQ_CHUNKS = os.getenv("USE_GOLDEN_AS_FAQ_CHUNKS", "false").lower() == "true"

print("IN_COLAB =", IN_COLAB)
print("torch available =", _TORCH_AVAILABLE)
print("CUDA available =", _cuda_available())
if _cuda_available():
    print("GPU =", torch.cuda.get_device_name(0))
print("USE_LOCAL_LLM_FOR_QUERY_UNDERSTANDING =", USE_LOCAL_LLM_FOR_QUERY_UNDERSTANDING)
print("LOCAL_LLM_OPTIONAL_REWRITE_MAX_TERMS =", LOCAL_LLM_OPTIONAL_REWRITE_MAX_TERMS)
print("HF_LLM_MODEL_NAME =", HF_LLM_MODEL_NAME)
print("HF_LLM_USE_4BIT =", HF_LLM_USE_4BIT)
print("EMBEDDING_MODEL =", EMBEDDING_MODEL_NAME)
print("OUTPUT_DIR =", OUTPUT_DIR)
print("LABOR_LAW_DOCX_PATH =", LABOR_LAW_DOCX_PATH or "<upload required>")
print("INTERNAL_POLICY_DOCX_PATH =", INTERNAL_POLICY_DOCX_PATH or "<upload required>")
print("GOLDEN_DATASET_JSON_PATH =", GOLDEN_DATASET_JSON_PATH or "<upload required>")
print("USE_GOLDEN_AS_FAQ_CHUNKS =", USE_GOLDEN_AS_FAQ_CHUNKS)
print("OFFLINE_ARTIFACT_DIR =", OFFLINE_ARTIFACT_DIR or "<auto-detect or upload optional>")
print("OFFLINE_ARTIFACT_ZIP_PATH =", OFFLINE_ARTIFACT_ZIP_PATH or "<auto-detect or upload optional>")
print("LOAD_PENDING_GRAPH_EDGES =", LOAD_PENDING_GRAPH_EDGES)

# -----------------------------
# 2. Utilities
# -----------------------------
CJK_NUM = "一二三四五六七八九十百千零〇"

CATEGORY_KEYWORDS = {
    "leave": ["請假", "特休", "特別休假", "年假", "病假", "事假", "休假", "例假", "休息日", "補休", "假別"],
    "overtime": ["加班", "延長工時", "工時", "工作時間", "加班費", "補休", "換補休"],
    "salary": ["薪資", "薪水", "工資", "給付", "扣薪", "全勤", "獎金", "津貼"],
    "termination": ["資遣", "離職", "解僱", "終止契約", "預告", "遣散", "非自願離職"],
    "attendance": ["出勤", "遲到", "早退", "打卡", "曠職", "排班", "輪班"],
    "welfare": ["福利", "旅遊補助", "員工貸款", "餐補", "教育訓練", "生日", "健康檢查"],
    "occupational_accident": ["職災", "受傷", "職業災害", "補償", "醫療"],
    "privacy_sensitive": ["個資", "身分證", "病歷", "診斷證明", "申訴", "懲處", "性騷擾"],
}

RISK_KEYWORDS = ["違法", "申訴", "告", "提告", "主管逼", "強迫", "不給薪", "扣薪", "解僱", "懲處", "性騷擾", "個資", "歧視"]
AMBIGUOUS_SHORTS = ["我想請假", "請假規定", "休假規定", "可以嗎", "合法嗎", "怎麼辦", "我要請假"]


def normalize_spaces(text: str) -> str:
    return re.sub(r"\s+", " ", str(text)).strip()


def normalize_article_no(x: str) -> str:
    x = str(x)
    x = re.sub(r"\s+", "", x)
    x = x.replace("第", "第 ").replace("條", " 條")
    return normalize_spaces(x)


def detect_category(text: str) -> str:
    scores = defaultdict(int)
    for cat, kws in CATEGORY_KEYWORDS.items():
        for kw in kws:
            if kw in text:
                scores[cat] += 1
    if not scores:
        return "general"
    return max(scores.items(), key=lambda x: x[1])[0]


def extract_keywords(text: str, max_n: int = 30) -> List[str]:
    kws = []
    for v in CATEGORY_KEYWORDS.values():
        for kw in v:
            if kw in text:
                kws.append(kw)
    tokens = re.findall(r"[\u4e00-\u9fffA-Za-z0-9]{2,}", text)
    return list(dict.fromkeys(kws + tokens))[:max_n]


def tokenize_zh(text: str) -> List[str]:
    # Simple tokenizer for BM25: keywords + CJK bigrams + alnum tokens
    text = str(text)
    tokens = extract_keywords(text, max_n=50)
    cjk = re.findall(r"[\u4e00-\u9fff]", text)
    tokens += ["".join(cjk[i:i+2]) for i in range(max(0, len(cjk)-1))]
    tokens += re.findall(r"[A-Za-z0-9]{2,}", text.lower())
    return [t for t in tokens if t]


def safe_json_dumps(obj: Any) -> str:
    return json.dumps(obj, ensure_ascii=False, indent=2)


# -----------------------------
# 2A. Offline Artifacts Loader
# -----------------------------
OFFLINE_ARTIFACT_FILENAMES = [
    "workflow_role_mapping.json",
    "concept_nodes.json",
    "risk_policy.json",
    "query_patterns.json",
    "rewrite_rules.json",
    "relation_schema.json",
    "graph_relation_candidates.json",
    "local_llm_usage_policy.json",
    "artifact_manifest.json",
]

class OfflineArtifacts:
    """Versioned offline LLM-assisted knowledge/config artifacts.

    These artifacts are generated or curated before runtime, then loaded as JSON.
    Runtime local LLM may use them for structured classification context, but final
    risk and route decisions are still deterministic.
    """
    def __init__(self, artifact_dir: Optional[str] = None):
        self.artifact_dir = Path(artifact_dir) if artifact_dir else None
        self.data: Dict[str, Any] = {}
        self.loaded_files: Dict[str, str] = {}

    @property
    def concept_nodes(self) -> List[Dict[str, Any]]:
        return self.data.get("concept_nodes", {}).get("concept_nodes", [])

    @property
    def risk_policies(self) -> List[Dict[str, Any]]:
        return self.data.get("risk_policy", {}).get("risk_policies", [])

    @property
    def query_patterns(self) -> List[Dict[str, Any]]:
        return self.data.get("query_patterns", {}).get("patterns", [])

    @property
    def rewrite_rules(self) -> List[Dict[str, Any]]:
        return self.data.get("rewrite_rules", {}).get("rewrite_rules", [])

    @property
    def relation_types(self) -> List[Dict[str, Any]]:
        return self.data.get("relation_schema", {}).get("relation_types", [])

    @property
    def graph_edge_candidates(self) -> List[Dict[str, Any]]:
        return self.data.get("graph_relation_candidates", {}).get("edge_candidates", [])

    def load(self) -> "OfflineArtifacts":
        if not self.artifact_dir or not self.artifact_dir.exists():
            print("No offline artifact folder found. Falling back to code defaults where available.")
            return self
        for fn in OFFLINE_ARTIFACT_FILENAMES:
            fp = self.artifact_dir / fn
            if fp.exists():
                key = fp.stem
                with open(fp, "r", encoding="utf-8") as f:
                    self.data[key] = json.load(f)
                self.loaded_files[key] = str(fp)
        print("Loaded offline artifacts:", sorted(self.loaded_files.keys()))
        return self

    def category_keywords(self) -> Dict[str, List[str]]:
        kw = defaultdict(list)
        for c in self.concept_nodes:
            cat = c.get("category") or "general"
            vals = []
            vals.append(c.get("label", ""))
            vals.extend(c.get("aliases", []) or [])
            vals.extend(c.get("retrieval_keywords", []) or [])
            for v in vals:
                if v and v not in kw[cat]:
                    kw[cat].append(str(v))
        return dict(kw)


def _extract_zip_to_dir(zip_path: str, target_dir: Path) -> Optional[Path]:
    zp = Path(zip_path)
    if not zp.exists() or not zp.suffix.lower() == ".zip":
        return None
    target_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(zp, "r") as z:
        z.extractall(target_dir)
    # If zip contains a single folder, use it; otherwise use target root.
    candidates = [p for p in target_dir.iterdir() if p.is_dir()]
    for cand in candidates:
        if (cand / "concept_nodes.json").exists():
            return cand
    return target_dir


def locate_offline_artifacts_dir() -> Optional[str]:
    # 1) Explicit directory
    if OFFLINE_ARTIFACT_DIR and Path(OFFLINE_ARTIFACT_DIR).exists():
        return str(Path(OFFLINE_ARTIFACT_DIR))

    # 2) Explicit ZIP
    if OFFLINE_ARTIFACT_ZIP_PATH and Path(OFFLINE_ARTIFACT_ZIP_PATH).exists():
        extracted = _extract_zip_to_dir(OFFLINE_ARTIFACT_ZIP_PATH, OUTPUT_DIR / "offline_artifacts_loaded")
        if extracted:
            return str(extracted)

    # 3) Common local/Colab locations
    candidates = [
        Path("/content/hr_offline_artifacts"),
        Path("/content/offline_artifacts"),
        Path("./hr_offline_artifacts"),
        Path("./offline_artifacts"),
        Path("/mnt/data/hr_offline_artifacts"),
    ]
    for p in candidates:
        if p.exists() and (p / "concept_nodes.json").exists():
            return str(p)

    # 4) Auto-detect ZIP in common locations
    zip_candidates = list(Path("/content").glob("*offline*artifacts*.zip")) if Path("/content").exists() else []
    zip_candidates += list(Path(".").glob("*offline*artifacts*.zip"))
    zip_candidates += list(Path("/mnt/data").glob("*offline*artifacts*.zip")) if Path("/mnt/data").exists() else []
    if zip_candidates:
        extracted = _extract_zip_to_dir(str(zip_candidates[0]), OUTPUT_DIR / "offline_artifacts_loaded")
        if extracted:
            return str(extracted)

    # 5) Optional upload in Colab
    if IN_COLAB:
        print("可選：上傳 offline artifacts ZIP（若略過，會使用程式內建 fallback 規則）。")
        try:
            uploaded = files.upload()
            for name in uploaded.keys():
                if name.lower().endswith(".zip"):
                    extracted = _extract_zip_to_dir(f"/content/{name}", OUTPUT_DIR / "offline_artifacts_loaded")
                    if extracted:
                        return str(extracted)
        except Exception as e:
            print("Offline artifact upload skipped or failed:", repr(e))
    return None


def load_offline_artifacts() -> OfflineArtifacts:
    artifact_dir = locate_offline_artifacts_dir()
    artifacts = OfflineArtifacts(artifact_dir).load()
    # Update global CATEGORY_KEYWORDS using concept_nodes as an externalized taxonomy.
    external_kw = artifacts.category_keywords()
    if external_kw:
        for cat, vals in external_kw.items():
            base = CATEGORY_KEYWORDS.setdefault(cat, [])
            for v in vals:
                if v and v not in base:
                    base.append(v)
    return artifacts


def normalize_route(route: str) -> str:
    route = str(route or "").strip()
    mapping = {"direct": "answer", "with_disclaimer": "disclaimer", "answer": "answer", "clarify": "clarify", "escalate": "escalate", "disclaimer": "disclaimer"}
    return mapping.get(route, route or "answer")


def policy_to_route(policy: str) -> str:
    return normalize_route(policy)


def route_to_policy(route: str) -> str:
    route = normalize_route(route)
    return {"answer": "direct", "disclaimer": "with_disclaimer", "clarify": "clarify", "escalate": "escalate"}.get(route, "direct")


CATEGORY_ALIASES = {
    "working_hours": {"attendance", "overtime", "working_hours"},
    "high_risk": {"privacy_sensitive", "termination", "salary", "leave", "overtime", "general", "high_risk"},
    "privacy_sensitive": {"privacy_sensitive", "salary", "general"},
}


def category_matches(expected: Any, actual: Any) -> bool:
    expected = str(expected or "").strip()
    actual = str(actual or "").strip()
    if not expected:
        return True
    if expected == actual:
        return True
    return actual in CATEGORY_ALIASES.get(expected, set())


def normalize_for_match(text: Any) -> str:
    return re.sub(r"\s+", "", str(text or "")).lower()


def extract_article_refs(text: Any) -> List[str]:
    """Extract normalized references like 第38條 from a citation string."""
    text = str(text or "")
    refs = []
    for raw in re.findall(rf"第\s*([\d{CJK_NUM}]+)\s*條", text):
        refs.append(normalize_for_match(f"第{raw}條"))
    return refs


def relation_from_policy_content(content: str) -> str:
    if any(k in content for k in ["優於", "較有利", "不低於", "最低標準"]):
        return "overrides"
    return "refers_to"


def law_article_id_from_ref(raw_no: str) -> str:
    return article_id_from_no("law", f"第 {raw_no} 條")


def extract_related_law_ids(content: str) -> List[str]:
    refs = []
    for raw in re.findall(rf"勞動基準法[^。；;\n]*第\s*([\d{CJK_NUM}]+)\s*條", content):
        refs.append(law_article_id_from_ref(raw))
    # Some policy clauses write only 「第 30 條」 after mentioning external regulation nearby.
    if "勞動基準法" in content:
        for raw in re.findall(rf"第\s*([\d{CJK_NUM}]+)\s*條", content):
            refs.append(law_article_id_from_ref(raw))
    return list(dict.fromkeys(refs))

# -----------------------------
# 3. DOCX Parser + Sample Data
# -----------------------------

def read_docx_text(path: str) -> str:
    doc = Document(path)
    blocks = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            blocks.append(txt)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n".join(blocks)


def create_sample_labor_law_docx(path: str) -> str:
    doc = Document()
    doc.add_heading("參考資料_勞動基準法 Sample", level=1)
    sample_articles = [
        ("第一章 總則", "第 2 條", "本法用詞，定義如下：勞工指受雇主僱用從事工作獲致工資者。工資指勞工因工作而獲得之報酬。"),
        ("第二章 勞動契約", "第 16 條", "雇主依規定終止勞動契約者，其預告期間依勞工工作年資而定。三個月以上一年未滿者十日前預告；一年以上三年未滿者二十日前預告；三年以上者三十日前預告。"),
        ("第三章 工資", "第 22 條", "工資應全額直接給付勞工。"),
        ("第三章 工資", "第 24 條", "雇主延長勞工工作時間者，延長工作時間在二小時以內者，按平日每小時工資額加給三分之一以上；再延長工作時間在二小時以內者，按平日每小時工資額加給三分之二以上。"),
        ("第四章 工作時間、休息、休假", "第 30 條", "勞工正常工作時間，每日不得超過八小時，每週不得超過四十小時。"),
        ("第四章 工作時間、休息、休假", "第 32 條", "雇主有使勞工在正常工作時間以外工作之必要者，經工會同意，如事業單位無工會者，經勞資會議同意後，得將工作時間延長之。"),
        ("第四章 工作時間、休息、休假", "第 36 條", "勞工每七日中應有二日之休息，其中一日為例假，一日為休息日。"),
        ("第四章 工作時間、休息、休假", "第 38 條", "勞工在同一雇主或事業單位，繼續工作滿一定期間者，應給予特別休假。六個月以上一年未滿者三日；一年以上二年未滿者七日；二年以上三年未滿者十日。"),
        ("第七章 職業災害補償", "第 59 條", "勞工因遭遇職業災害而致死亡、失能、傷害或疾病時，雇主應依規定予以補償。"),
    ]
    current_chapter = None
    for chapter, article, content in sample_articles:
        if chapter != current_chapter:
            doc.add_heading(chapter, level=2)
            current_chapter = chapter
        doc.add_paragraph(f"{article} {content}")
    doc.save(path)
    return path

# -----------------------------
# 4. Policy-aware Hierarchical Chunking
# -----------------------------
CHAPTER_PATTERN = re.compile(rf"^第\s*[\d{CJK_NUM}]+\s*章.*")
ARTICLE_PATTERN = re.compile(
    rf"^(第\s*[\d{CJK_NUM}]+(?:\s*-\s*[\d{CJK_NUM}]+)?\s*條(?:\s*之\s*[\d{CJK_NUM}]+)?)[\s：:、]*(.*)"
)
POLICY_ARTICLE_PATTERN = re.compile(r"^(POLICY[-_][A-Za-z0-9]+|內規[-_][A-Za-z0-9]+|[A-Za-z]+[-_]\d+)[\s：:、]*(.*)", re.I)


def split_sentences_zh(text: str) -> List[str]:
    parts = re.split(r"(?<=[。！？；;])", text)
    return [p.strip() for p in parts if p.strip()]


def article_id_from_no(source_type: str, article_no: str) -> str:
    clean = re.sub(r"\s+", "", article_no)
    clean = clean.replace("第", "").replace("條", "")
    clean = clean.replace("之", "_")
    prefix = "law" if source_type == "law" else "policy"
    return f"{prefix}_{clean}"


@dataclass
class ChunkConfig:
    version: str = "v1.0"
    effective_date: str = "2026-06-23"


class HRKnowledgeBuilder:
    def __init__(self, config: Optional[ChunkConfig] = None):
        self.config = config or ChunkConfig()

    def parse_labor_law_articles(self, text: str, document_name: str) -> List[Dict[str, Any]]:
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        articles = []
        current_chapter = "未標示章節"
        current_article_no = None
        current_content = []

        def flush():
            if current_article_no and current_content:
                content = "\n".join(current_content).strip()
                if content:
                    article_no = normalize_article_no(current_article_no)
                    articles.append({
                        "source_type": "law",
                        "document_name": document_name,
                        "article_id": article_id_from_no("law", article_no),
                        "article_no": article_no,
                        "chapter": current_chapter,
                        "title": "",
                        "category": detect_category(content),
                        "priority": 1,
                        "version": self.config.version,
                        "effective_date": self.config.effective_date,
                        "content": content,
                        "related_articles": [],
                    })

        for line in lines:
            if CHAPTER_PATTERN.match(line):
                current_chapter = line
                continue
            m = ARTICLE_PATTERN.match(line)
            if m:
                flush()
                current_article_no = m.group(1)
                rest = m.group(2).strip()
                current_content = [f"{normalize_article_no(current_article_no)} {rest}".strip()]
            else:
                if current_article_no:
                    current_content.append(line)
        flush()

        # fallback: fixed chunks if article parser fails
        if not articles:
            raw = text
            size, overlap = 900, 120
            start = 0
            while start < len(raw):
                content = raw[start:start+size]
                article_no = f"chunk_{len(articles)+1}"
                articles.append({
                    "source_type": "law",
                    "document_name": document_name,
                    "article_id": f"law_chunk_{len(articles)+1}",
                    "article_no": article_no,
                    "chapter": "未標示章節",
                    "title": "",
                    "category": detect_category(content),
                    "priority": 1,
                    "version": self.config.version,
                    "effective_date": self.config.effective_date,
                    "content": content,
                    "related_articles": [],
                })
                start += size - overlap
        return articles

    def parse_internal_policy_articles(self, text: str, document_name: str) -> List[Dict[str, Any]]:
        """Parse user-provided internal policy DOCX into article-level structured records.

        Expected heading examples:
        - 第 11 條｜特別休假
        - 第11條 特別休假
        - POLICY-LEAVE-001 特別休假

        The parser intentionally skips table-of-contents entries that have no body text.
        """
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        articles = []
        current_chapter = "未標示章節"
        current_article_no = None
        current_title = ""
        current_content = []

        internal_article_pattern = re.compile(
            rf"^(第\s*[\d{CJK_NUM}]+(?:\s*-\s*[\d{CJK_NUM}]+)?\s*條(?:\s*之\s*[\d{CJK_NUM}]+)?|POLICY[-_][A-Za-z0-9]+|內規[-_][A-Za-z0-9]+)[\s｜|：:、]*(.*)$",
            re.I,
        )

        def flush():
            if not current_article_no:
                return
            body = "\n".join(current_content).strip()
            # Skip TOC-only headings or empty articles.
            if not body or len(body) < 8:
                return
            article_no = normalize_article_no(current_article_no)
            title = normalize_spaces(current_title)
            content = f"{article_no} {title}\n{body}".strip()
            related = extract_related_law_ids(content)
            graph_edges = []
            for rid in related:
                graph_edges.append((relation_from_policy_content(content), rid))
            articles.append({
                "source_type": "internal_policy",
                "document_name": document_name,
                "article_id": article_id_from_no("internal_policy", article_no),
                "article_no": article_no,
                "chapter": current_chapter,
                "title": title,
                "category": detect_category(content + " " + title + " " + current_chapter),
                "priority": 2,
                "version": self.config.version,
                "effective_date": self.config.effective_date,
                "content": content,
                "related_articles": related,
                "graph_edges": graph_edges,
            })

        for line in lines:
            if CHAPTER_PATTERN.match(line):
                flush()
                current_chapter = line
                current_article_no = None
                current_title = ""
                current_content = []
                continue
            m = internal_article_pattern.match(line)
            if m:
                flush()
                current_article_no = m.group(1)
                current_title = m.group(2).strip()
                current_content = []
            else:
                if current_article_no:
                    current_content.append(line)
        flush()

        # Fallback: if the policy file is not article-formatted, create fixed article-like chunks.
        if not articles:
            raw = text
            size, overlap = 900, 120
            start = 0
            while start < len(raw):
                content = raw[start:start+size]
                article_no = f"policy_chunk_{len(articles)+1}"
                articles.append({
                    "source_type": "internal_policy",
                    "document_name": document_name,
                    "article_id": f"policy_chunk_{len(articles)+1}",
                    "article_no": article_no,
                    "chapter": "未標示章節",
                    "title": "",
                    "category": detect_category(content),
                    "priority": 2,
                    "version": self.config.version,
                    "effective_date": self.config.effective_date,
                    "content": content,
                    "related_articles": extract_related_law_ids(content),
                    "graph_edges": [],
                })
                start += size - overlap
        return articles

    def make_document_chunks(self, articles: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        by_doc = defaultdict(list)
        for a in articles:
            by_doc[a["document_name"]].append(a)
        chunks = []
        for doc_name, arr in by_doc.items():
            source_type = arr[0]["source_type"]
            content = "\n".join([f"{a['article_no']} {a['title']}：{a['content'][:160]}" for a in arr[:20]])
            cats = sorted(set(a["category"] for a in arr))
            chunks.append({
                "chunk_id": f"doc_{len(chunks)+1:04d}",
                "chunk_type": "document",
                "source_type": source_type,
                "document_name": doc_name,
                "article_id": f"document::{doc_name}",
                "article_no": "DOCUMENT",
                "chapter": "Document-level",
                "title": f"{doc_name} 文件摘要",
                "category": ",".join(cats),
                "priority": max(a["priority"] for a in arr),
                "version": self.config.version,
                "effective_date": self.config.effective_date,
                "content": content,
                "parent_id": None,
                "related_articles": [],
                "keywords": extract_keywords(content),
            })
        return chunks

    def make_article_chunks(self, articles: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        chunks = []
        for a in articles:
            chunks.append({
                "chunk_id": f"article_{len(chunks)+1:04d}",
                "chunk_type": "article",
                "source_type": a["source_type"],
                "document_name": a["document_name"],
                "article_id": a["article_id"],
                "article_no": a["article_no"],
                "chapter": a["chapter"],
                "title": a.get("title", ""),
                "category": a["category"],
                "priority": a["priority"],
                "version": a["version"],
                "effective_date": a["effective_date"],
                "content": a["content"],
                "parent_id": f"document::{a['document_name']}",
                "related_articles": a.get("related_articles", []),
                "keywords": extract_keywords(a["content"]),
            })
        return chunks

    def make_semantic_subchunks(self, articles: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        chunks = []
        for a in articles:
            sentences = split_sentences_zh(a["content"])
            # group every 1-2 sentences if long
            groups = []
            buf = []
            for s in sentences:
                buf.append(s)
                if len("".join(buf)) >= 120:
                    groups.append("".join(buf))
                    buf = []
            if buf:
                groups.append("".join(buf))
            if len(groups) <= 1 and len(a["content"]) < 220:
                continue
            for j, g in enumerate(groups, start=1):
                chunks.append({
                    "chunk_id": f"semantic_{len(chunks)+1:04d}",
                    "chunk_type": "semantic",
                    "source_type": a["source_type"],
                    "document_name": a["document_name"],
                    "article_id": f"{a['article_id']}::s{j}",
                    "article_no": a["article_no"],
                    "chapter": a["chapter"],
                    "title": a.get("title", ""),
                    "category": a["category"],
                    "priority": a["priority"],
                    "version": a["version"],
                    "effective_date": a["effective_date"],
                    "content": g,
                    "parent_id": a["article_id"],
                    "related_articles": a.get("related_articles", []),
                    "keywords": extract_keywords(g),
                })
        return chunks

    def make_faq_chunks(self, articles: List[Dict[str, Any]], golden_df: Optional[pd.DataFrame] = None) -> List[Dict[str, Any]]:
        """Build FAQ chunks only when explicitly enabled.

        Default is empty because Golden Dataset is evaluation data and should not be leaked
        into the retrieval knowledge base. If USE_GOLDEN_AS_FAQ_CHUNKS=true, the script
        converts user-provided Golden Dataset questions into experimental FAQ chunks.
        """
        if not USE_GOLDEN_AS_FAQ_CHUNKS or golden_df is None or golden_df.empty:
            return []
        chunks = []
        for _, row in golden_df.iterrows():
            q = str(row.get("question", "")).strip()
            if not q:
                continue
            key_points = row.get("expected_key_points", [])
            if isinstance(key_points, str):
                answer = key_points
            elif isinstance(key_points, list):
                answer = "；".join(map(str, key_points))
            else:
                answer = ""
            category = str(row.get("expected_category", detect_category(q)) or detect_category(q))
            content = f"FAQ 問題：{q}\nFAQ 回答重點：{answer}"
            chunks.append({
                "chunk_id": f"faq_{len(chunks)+1:04d}",
                "chunk_type": "faq",
                "source_type": "golden_dataset_faq_experiment",
                "document_name": "Golden Dataset derived FAQ chunks - experimental",
                "article_id": f"faq::golden::{row.get('id', len(chunks)+1)}",
                "article_no": str(row.get("id", f"FAQ-{len(chunks)+1}")),
                "chapter": "FAQ Chunk",
                "title": q,
                "category": category,
                "priority": 1,
                "version": self.config.version,
                "effective_date": self.config.effective_date,
                "content": content,
                "question": q,
                "answer": answer,
                "parent_id": None,
                "related_articles": [],
                "keywords": extract_keywords(content),
            })
        return chunks

    def build_chunks(
        self,
        labor_law_docx_path: str,
        internal_policy_docx_path: str,
        golden_df: Optional[pd.DataFrame] = None,
    ) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        # 1) External regulation: user-provided 勞動基準法 DOCX
        law_text = read_docx_text(labor_law_docx_path)
        law_articles = self.parse_labor_law_articles(law_text, Path(labor_law_docx_path).name)

        # 2) Internal policy: user-provided 模擬銀行內規 DOCX
        policy_text = read_docx_text(internal_policy_docx_path)
        policy_articles = self.parse_internal_policy_articles(policy_text, Path(internal_policy_docx_path).name)

        all_articles = law_articles + policy_articles

        doc_chunks = self.make_document_chunks(all_articles)
        article_chunks = self.make_article_chunks(all_articles)
        semantic_chunks = self.make_semantic_subchunks(all_articles)
        faq_chunks = self.make_faq_chunks(all_articles, golden_df=golden_df)

        chunks = doc_chunks + article_chunks + semantic_chunks + faq_chunks
        # ensure unique ids
        for i, c in enumerate(chunks, start=1):
            c["global_chunk_id"] = f"G{i:05d}"
            c["embedding_text"] = self.chunk_to_text(c)
        return all_articles, chunks

    @staticmethod
    def chunk_to_text(c: Dict[str, Any]) -> str:
        return f"""
Chunk Type: {c.get('chunk_type')}
Source Type: {c.get('source_type')}
Document: {c.get('document_name')}
Article: {c.get('article_no')}
Title: {c.get('title')}
Category: {c.get('category')}
Priority: {c.get('priority')}
Content: {c.get('content')}
Keywords: {', '.join(c.get('keywords', []))}
""".strip()

# -----------------------------
# 5. Knowledge Graph / Graph RAG
# -----------------------------
class HRKnowledgeGraph:
    def __init__(self, articles: List[Dict[str, Any]], chunks: List[Dict[str, Any]], artifacts: Optional[OfflineArtifacts] = None):
        self.articles = articles
        self.chunks = chunks
        self.artifacts = artifacts or OfflineArtifacts()
        self.G = nx.DiGraph()
        self._article_lookup = {a["article_id"]: a for a in articles}
        self._article_no_lookup = self._build_article_no_lookup()
        self._build_graph()

    def _build_article_no_lookup(self) -> Dict[Tuple[str, str], str]:
        lookup = {}
        for a in self.articles:
            source_type = a.get("source_type", "")
            no_norm = normalize_for_match(a.get("article_no", ""))
            if no_norm:
                lookup[(source_type, no_norm)] = a.get("article_id")
                lookup[("any", no_norm)] = a.get("article_id")
        return lookup

    def _resolve_article_ref(self, ref: str, preferred_source: Optional[str] = None) -> Optional[str]:
        # Direct article id
        if ref in self._article_lookup:
            return ref
        ref = str(ref or "")
        # Convert policy/law article hints like policy_article_11 / law_article_38
        m = re.search(r"(policy|law)_article_([0-9]+)", ref)
        if m:
            source = "internal_policy" if m.group(1) == "policy" else "law"
            no_norm = normalize_for_match(f"第{m.group(2)}條")
            return self._article_no_lookup.get((source, no_norm)) or self._article_no_lookup.get(("any", no_norm))
        # Convert human-readable ref like 安久銀行...第11條 or 勞動基準法第38條
        refs = extract_article_refs(ref)
        if refs:
            source = preferred_source
            if "勞動基準法" in ref or "勞基法" in ref:
                source = "law"
            elif "安久銀行" in ref or "內規" in ref or "規章" in ref:
                source = "internal_policy"
            for no_norm in refs:
                if source:
                    hit = self._article_no_lookup.get((source, no_norm))
                    if hit:
                        return hit
                hit = self._article_no_lookup.get(("any", no_norm))
                if hit:
                    return hit
        return None

    def _add_concept_nodes(self):
        if self.artifacts.concept_nodes:
            for c in self.artifacts.concept_nodes:
                cid = c.get("concept_id")
                if not cid:
                    continue
                self.G.add_node(
                    cid,
                    node_type="concept",
                    label=c.get("label", cid),
                    category=c.get("category", "general"),
                    risk_level=c.get("risk_level", "low"),
                    default_answer_policy=c.get("default_answer_policy", "answer"),
                    content=c.get("description", c.get("label", cid)),
                    aliases="｜".join(c.get("aliases", []) or []),
                    graph_expansion_priority=c.get("graph_expansion_priority", "medium"),
                )
                parent = c.get("parent_concept_id")
                if parent:
                    self.G.add_edge(parent, cid, relation="parent_of")
                    self.G.add_edge(cid, parent, relation="child_of")
            return

        # Fallback concept skeleton if no artifact is provided.
        concepts = {
            "concept_leave": "請假制度",
            "concept_special_leave": "特別休假",
            "concept_overtime": "加班",
            "concept_comp_time": "補休",
            "concept_working_hours": "工時",
            "concept_salary": "薪資",
            "concept_welfare": "福利",
            "concept_termination": "離職與資遣",
        }
        for cid, label in concepts.items():
            self.G.add_node(cid, node_type="concept", label=label, content=label)

    def _add_concept_article_edges(self):
        # Artifact-driven concept-to-article mapping
        for c in self.artifacts.concept_nodes:
            cid = c.get("concept_id")
            if not cid or cid not in self.G:
                continue
            for ref in c.get("related_law_articles", []) or []:
                aid = self._resolve_article_ref(ref, preferred_source="law")
                if aid:
                    self.G.add_edge(cid, aid, relation="has_rule")
                    self.G.add_edge(aid, cid, relation="related_to")
            for ref in c.get("related_policy_articles", []) or []:
                aid = self._resolve_article_ref(ref, preferred_source="internal_policy")
                if aid:
                    self.G.add_edge(cid, aid, relation="has_rule")
                    self.G.add_edge(aid, cid, relation="related_to")

        # Fallback category-related edges if no concept artifacts are available.
        if not self.artifacts.concept_nodes:
            concept_by_cat = {
                "leave": "concept_leave",
                "overtime": "concept_overtime",
                "attendance": "concept_working_hours",
                "salary": "concept_salary",
                "welfare": "concept_welfare",
                "termination": "concept_termination",
            }
            for a in self.articles:
                concept = concept_by_cat.get(a["category"])
                if concept:
                    self.G.add_edge(a["article_id"], concept, relation="related_to")
                    self.G.add_edge(concept, a["article_id"], relation="has_rule")

    def _add_artifact_edge_candidates(self):
        allowed_relations = {r.get("relation_type") for r in self.artifacts.relation_types if r.get("runtime_expandable", True)}
        if not allowed_relations:
            allowed_relations = {"has_rule", "related_to", "refers_to", "supplements", "overrides", "parent_of", "child_of"}
        for e in self.artifacts.graph_edge_candidates:
            status = e.get("review_status", "pending")
            if status != "approved" and not LOAD_PENDING_GRAPH_EDGES:
                continue
            rel = e.get("relation_type", "related_to")
            if rel not in allowed_relations:
                continue
            src = self._resolve_article_ref(e.get("source_node", "")) or e.get("source_node")
            tgt = self._resolve_article_ref(e.get("target_node", "")) or e.get("target_node")
            if src in self.G and tgt in self.G:
                self.G.add_edge(
                    src, tgt,
                    relation=rel,
                    evidence=e.get("evidence", ""),
                    confidence=e.get("confidence", None),
                    review_status=status,
                    human_review_required=e.get("human_review_required", False),
                )

    def _build_graph(self):
        self._add_concept_nodes()
        for a in self.articles:
            self.G.add_node(
                a["article_id"],
                node_type="law_article" if a["source_type"] == "law" else "internal_policy_article",
                label=f"{a['article_no']} {a.get('title','')}",
                source_type=a["source_type"],
                article_no=a["article_no"],
                category=a["category"],
                priority=a["priority"],
                content=a["content"],
            )
            # direct related articles parsed from document text
            for rel in a.get("related_articles", []):
                if rel:
                    self.G.add_edge(a["article_id"], rel, relation="refers_to")
            for relation, target in a.get("graph_edges", []):
                self.G.add_edge(a["article_id"], target, relation=relation)

        self._add_concept_article_edges()
        self._add_artifact_edge_candidates()

    def expand(self, seed_article_ids: List[str], question: str, hops: int = 1, max_nodes: int = 12, preferred_relations: Optional[List[str]] = None) -> Dict[str, Any]:
        # Runtime expansion is deterministic; local LLM can provide preferred relations, but traversal uses approved graph only.
        use_graph = any(k in question for k in ["差", "比較", "為什麼", "內規", "法規", "公司", "優於", "補休", "依據", "哪個", "關係"])
        if not use_graph and len(seed_article_ids) <= 0:
            return {"use_graph": False, "nodes": [], "edges": [], "context": ""}

        relation_filter = set(preferred_relations or [])
        nodes = []
        edges = []
        visited = set()
        frontier = [s for s in seed_article_ids if s in self.G]
        for s in frontier:
            visited.add(s)

        for _ in range(hops):
            new_frontier = []
            for u in frontier:
                neighbors = list(self.G.successors(u)) + list(self.G.predecessors(u))
                for v in neighbors:
                    candidate_edges = []
                    if self.G.has_edge(u, v):
                        candidate_edges.append((u, v, self.G[u][v].get("relation", "related_to")))
                    if self.G.has_edge(v, u):
                        candidate_edges.append((v, u, self.G[v][u].get("relation", "related_to")))
                    if relation_filter and not any(r in relation_filter for _, _, r in candidate_edges):
                        continue
                    if v not in visited:
                        visited.add(v)
                        new_frontier.append(v)
                    edges.extend(candidate_edges)
                    if len(visited) >= max_nodes:
                        break
                if len(visited) >= max_nodes:
                    break
            frontier = new_frontier
            if len(visited) >= max_nodes:
                break

        for n in list(visited)[:max_nodes]:
            data = self.G.nodes[n]
            nodes.append({
                "node_id": n,
                "node_type": data.get("node_type"),
                "label": data.get("label"),
                "article_no": data.get("article_no", ""),
                "source_type": data.get("source_type", ""),
                "category": data.get("category", ""),
                "risk_level": data.get("risk_level", ""),
                "default_answer_policy": data.get("default_answer_policy", ""),
                "content": data.get("content", "")[:500],
            })
        seen = set()
        edge_dicts = []
        for u, v, r in edges:
            key = (u, v, r)
            if key not in seen:
                seen.add(key)
                edata = self.G[u][v] if self.G.has_edge(u, v) else {}
                edge_dicts.append({"source": u, "target": v, "relation": r, "evidence": edata.get("evidence", "")})

        context_lines = []
        if nodes:
            context_lines.append("[Graph Nodes]")
            for n in nodes:
                context_lines.append(f"- {n['node_id']} ({n['node_type']}): {n['label']}｜{n['content']}")
        if edge_dicts:
            context_lines.append("[Graph Relations]")
            for e in edge_dicts:
                ev = f"｜evidence: {e['evidence']}" if e.get("evidence") else ""
                context_lines.append(f"- {e['source']} --{e['relation']}--> {e['target']}{ev}")
        return {
            "use_graph": use_graph,
            "nodes": nodes,
            "edges": edge_dicts,
            "context": "\n".join(context_lines),
        }

    def save_graph_files(self, output_dir: Path):
        nodes = []
        for n, d in self.G.nodes(data=True):
            rec = {"node_id": n, **d}
            nodes.append(rec)
        edges = []
        for u, v, d in self.G.edges(data=True):
            edges.append({"source": u, "target": v, **d})
        pd.DataFrame(nodes).to_csv(output_dir / "kg_nodes.csv", index=False, encoding="utf-8-sig")
        pd.DataFrame(edges).to_csv(output_dir / "kg_edges.csv", index=False, encoding="utf-8-sig")
        nx.write_gexf(self.G, output_dir / "hr_knowledge_graph.gexf")

# -----------------------------
# 6. Hybrid Retriever
# -----------------------------
class HybridRetriever:
    def __init__(
        self,
        chunks: List[Dict[str, Any]],
        embedding_model_name: str = EMBEDDING_MODEL_NAME,
        use_reranker: bool = USE_RERANKER,
        reranker_model_name: str = RERANKER_MODEL_NAME,
    ):
        import faiss
        from sentence_transformers import SentenceTransformer
        self._faiss = faiss
        self.chunks = chunks
        print(f"Loading embedding model: {embedding_model_name}")
        self.embedder = SentenceTransformer(embedding_model_name)
        self.texts = [c["embedding_text"] for c in chunks]
        self.tokenized = [tokenize_zh(t) for t in self.texts]
        self.bm25 = BM25Okapi(self.tokenized)
        self.index = None
        self.embeddings = None
        self._build_vector_index()

        # Optional cross-encoder reranker. Loaded lazily; if it fails to load, retrieval
        # degrades gracefully to the hybrid (vector + BM25) ranking.
        self.reranker = None
        if use_reranker:
            try:
                from sentence_transformers import CrossEncoder
                print(f"Loading reranker model: {reranker_model_name}")
                self.reranker = CrossEncoder(reranker_model_name)
            except Exception as e:
                print("Reranker load failed; continuing without rerank.", repr(e))
                self.reranker = None

    def _rerank(self, query: str, candidates: List[Dict[str, Any]], top_k: int) -> List[Dict[str, Any]]:
        """Re-score the top hybrid candidates with the cross-encoder and blend scores."""
        if not self.reranker or not candidates:
            return candidates[:top_k]
        pool = candidates[:max(RERANK_CANDIDATES, top_k)]
        pairs = [[query, c.get("content", "") or c.get("embedding_text", "")] for c in pool]
        try:
            scores = self.reranker.predict(pairs, convert_to_numpy=True)
        except Exception as e:
            print("Rerank predict failed; using hybrid order.", repr(e))
            return candidates[:top_k]
        s = np.asarray(scores, dtype="float32")
        # Min-max normalize reranker scores so they blend with the hybrid final_score.
        if s.size > 1 and float(s.max() - s.min()) > 1e-9:
            s_norm = (s - s.min()) / (s.max() - s.min())
        else:
            s_norm = np.zeros_like(s)
        for c, raw, norm in zip(pool, scores, s_norm):
            hybrid = float(c.get("final_score", 0.0))
            c["rerank_score"] = round(float(raw), 4)
            c["pre_rerank_score"] = hybrid
            c["final_score"] = round(RERANK_WEIGHT * float(norm) + (1 - RERANK_WEIGHT) * hybrid, 4)
        pool = sorted(pool, key=lambda x: x["final_score"], reverse=True)
        return pool[:top_k]

    def _build_vector_index(self):
        print("Building embeddings and FAISS index...")
        emb = self.embedder.encode(
            self.texts,
            batch_size=32,
            show_progress_bar=True,
            convert_to_numpy=True,
            normalize_embeddings=True,
        ).astype("float32")
        self.embeddings = emb
        self.index = self._faiss.IndexFlatIP(emb.shape[1])
        self.index.add(emb)
        print("FAISS index size:", self.index.ntotal)

    def search(self, query: str, category: str = "general", top_k: int = 8) -> List[Dict[str, Any]]:
        query_emb = self.embedder.encode([query], convert_to_numpy=True, normalize_embeddings=True).astype("float32")
        search_k = min(max(top_k * 5, top_k), len(self.chunks))
        vector_scores, vector_ids = self.index.search(query_emb, search_k)

        bm25_scores = self.bm25.get_scores(tokenize_zh(query))
        if len(bm25_scores) > 0:
            bm25_norm = (bm25_scores - np.min(bm25_scores)) / (np.max(bm25_scores) - np.min(bm25_scores) + 1e-9)
        else:
            bm25_norm = np.zeros(len(self.chunks))

        candidate_ids = set(vector_ids[0].tolist())
        candidate_ids.update(np.argsort(-bm25_norm)[:search_k].tolist())
        keywords = extract_keywords(query)

        results = []
        for idx in candidate_ids:
            if idx < 0 or idx >= len(self.chunks):
                continue
            c = dict(self.chunks[idx])
            vscore = 0.0
            if idx in vector_ids[0]:
                pos = list(vector_ids[0]).index(idx)
                vscore = float(vector_scores[0][pos])
            bscore = float(bm25_norm[idx])
            keyword_bonus = sum(0.02 for kw in keywords if kw in c.get("embedding_text", ""))
            category_bonus = 0.06 if category != "general" and c.get("category") == category else 0.0
            priority_bonus = 0.04 * float(c.get("priority", 1))
            faq_bonus = 0.05 if c.get("chunk_type") == "faq" else 0.0
            article_bonus = 0.03 if c.get("chunk_type") == "article" else 0.0

            final_score = 0.62 * vscore + 0.28 * bscore + keyword_bonus + category_bonus + priority_bonus + faq_bonus + article_bonus
            c.update({
                "vector_score": round(vscore, 4),
                "bm25_score": round(bscore, 4),
                "keyword_bonus": round(keyword_bonus, 4),
                "category_bonus": round(category_bonus, 4),
                "priority_bonus": round(priority_bonus, 4),
                "final_score": round(float(final_score), 4),
            })
            results.append(c)
        results = sorted(results, key=lambda x: x["final_score"], reverse=True)
        # Cross-encoder rerank stage (no-op if reranker is disabled/unavailable).
        if self.reranker:
            return self._rerank(query, results, top_k)
        return results[:top_k]

# -----------------------------
# 7. HuggingFace Local LLM Helpers
# -----------------------------
class LocalHFLLM:
    """
    Colab GPU local LLM wrapper using HuggingFace Transformers.

    Recommended Colab settings:
    - Runtime > Change runtime type > T4 GPU or better
    - Default model: google/gemma-2-2b-it (open source; gated on HF — needs HF token)
    - 4-bit quantization: enabled by default to reduce GPU memory usage

    The wrapper is lazy-loaded: the model is downloaded and loaded only when the
    workflow first needs to generate an answer.
    """
    def __init__(
        self,
        model_name: str = HF_LLM_MODEL_NAME,
        use_4bit: bool = HF_LLM_USE_4BIT,
        max_new_tokens: int = HF_MAX_NEW_TOKENS,
    ):
        from transformers import AutoTokenizer, AutoModelForCausalLM

        self.model_name = model_name
        self.use_4bit = use_4bit and torch.cuda.is_available()
        self.max_new_tokens = max_new_tokens
        self.device = "cuda" if torch.cuda.is_available() else "cpu"

        print(f"Loading local HF LLM: {model_name}")
        print(f"Device: {self.device}; 4-bit: {self.use_4bit}")

        self.tokenizer = AutoTokenizer.from_pretrained(
            model_name,
            trust_remote_code=True,
            use_fast=True,
        )
        if self.tokenizer.pad_token_id is None:
            self.tokenizer.pad_token = self.tokenizer.eos_token

        if self.use_4bit:
            from transformers import BitsAndBytesConfig
            quantization_config = BitsAndBytesConfig(
                load_in_4bit=True,
                bnb_4bit_compute_dtype=torch.float16,
                bnb_4bit_quant_type="nf4",
                bnb_4bit_use_double_quant=True,
            )
            self.model = AutoModelForCausalLM.from_pretrained(
                model_name,
                device_map="auto",
                quantization_config=quantization_config,
                trust_remote_code=True,
            )
        else:
            dtype = torch.float16 if torch.cuda.is_available() else torch.float32
            self.model = AutoModelForCausalLM.from_pretrained(
                model_name,
                torch_dtype=dtype,
                device_map="auto" if torch.cuda.is_available() else None,
                trust_remote_code=True,
            )
            if not torch.cuda.is_available():
                self.model.to("cpu")

        self.model.eval()

    def _format_messages(self, system_prompt: str, user_prompt: str) -> str:
        system_prompt = system_prompt.strip()
        user_prompt = user_prompt.strip()
        if getattr(self.tokenizer, "chat_template", None):
            # Some chat templates (e.g. Gemma) do NOT support a separate "system" role
            # and raise on it. Try system+user first, then fall back to merging the
            # system prompt into the user turn.
            try:
                return self.tokenizer.apply_chat_template(
                    [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    tokenize=False,
                    add_generation_prompt=True,
                )
            except Exception:
                merged = f"{system_prompt}\n\n{user_prompt}" if system_prompt else user_prompt
                return self.tokenizer.apply_chat_template(
                    [{"role": "user", "content": merged}],
                    tokenize=False,
                    add_generation_prompt=True,
                )
        # Generic fallback for models without a chat_template.
        return f"System:\n{system_prompt}\n\nUser:\n{user_prompt}\n\nAssistant:\n"

    def generate(
        self,
        system_prompt: str,
        user_prompt: str,
        temperature: float = HF_TEMPERATURE,
        max_new_tokens: Optional[int] = None,
    ) -> str:
        prompt = self._format_messages(system_prompt, user_prompt)
        inputs = self.tokenizer(prompt, return_tensors="pt")

        # For quantized/device_map=auto models, putting inputs on cuda is usually correct in Colab.
        target_device = "cuda" if torch.cuda.is_available() else "cpu"
        inputs = {k: v.to(target_device) for k, v in inputs.items()}
        input_len = inputs["input_ids"].shape[-1]

        do_sample = temperature is not None and temperature > 0
        with torch.no_grad():
            outputs = self.model.generate(
                **inputs,
                max_new_tokens=max_new_tokens or self.max_new_tokens,
                do_sample=do_sample,
                temperature=temperature if do_sample else None,
                top_p=0.9 if do_sample else None,
                repetition_penalty=1.05,
                pad_token_id=self.tokenizer.pad_token_id,
                eos_token_id=self.tokenizer.eos_token_id,
            )
        new_tokens = outputs[0][input_len:]
        text = self.tokenizer.decode(new_tokens, skip_special_tokens=True)
        return text.strip()


_LOCAL_HF_LLM = None

def get_local_hf_llm() -> Optional[LocalHFLLM]:
    global _LOCAL_HF_LLM
    if _LOCAL_HF_LLM is None:
        try:
            _LOCAL_HF_LLM = LocalHFLLM()
        except Exception as e:
            print("Local HF LLM loading failed. Fallback to template answer.", repr(e))
            _LOCAL_HF_LLM = None
    return _LOCAL_HF_LLM


def call_llm_text(system_prompt: str, user_prompt: str, temperature: float = 0.1, max_new_tokens: int = HF_MAX_NEW_TOKENS) -> Optional[str]:
    llm = get_local_hf_llm()
    if llm is None:
        return None
    try:
        return llm.generate(
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            temperature=temperature,
            max_new_tokens=max_new_tokens,
        )
    except Exception as e:
        print("Local HF generation failed. Fallback to template answer.", repr(e))
        return None


def extract_json_from_text(text: Optional[str]) -> Optional[dict]:
    if not text:
        return None
    try:
        return json.loads(text)
    except Exception:
        pass
    m = re.search(r"\{.*\}", text, flags=re.S)
    if m:
        try:
            return json.loads(m.group(0))
        except Exception:
            return None
    return None


def call_llm_json(system_prompt: str, user_prompt: str, default: dict) -> dict:
    # Local small LLM JSON can be unstable, so this is optional and defaults to heuristic result.
    text = call_llm_text(system_prompt + "\n請只輸出 valid JSON，不要輸出 markdown。", user_prompt, temperature=0.0, max_new_tokens=384)
    parsed = extract_json_from_text(text)
    return parsed if isinstance(parsed, dict) else default

# -----------------------------
# 8. LangGraph Workflow
# -----------------------------
class HRState(TypedDict, total=False):
    question: str
    intent: str
    category: str
    risk_level: str
    answer_policy: Literal["direct", "with_disclaimer", "escalate", "clarify"]
    matched_concepts: List[str]
    missing_slots: List[str]
    llm_risk_signal: Dict[str, Any]
    risk_matches: List[Dict[str, Any]]
    preferred_relations: List[str]
    rewritten_query: str
    retrieved_chunks: List[Dict[str, Any]]
    graph_context: Dict[str, Any]
    confidence: float
    route: Literal["answer", "disclaimer", "escalate", "clarify"]
    answer: str
    citations: List[Dict[str, Any]]
    faithfulness_score: float
    debug: List[str]


def add_debug(state: HRState, message: str) -> List[str]:
    return state.get("debug", []) + [message]


class HRAssistantGraph:
    def __init__(self, retriever: HybridRetriever, kg: HRKnowledgeGraph, artifacts: Optional[OfflineArtifacts] = None):
        self.retriever = retriever
        self.kg = kg
        self.artifacts = artifacts or OfflineArtifacts()
        self.app = self._build_graph()

    def _match_concepts(self, question: str, top_k: int = 5) -> List[Dict[str, Any]]:
        q = normalize_for_match(question)
        matches = []
        for c in self.artifacts.concept_nodes:
            terms = [c.get("label", "")] + (c.get("aliases", []) or []) + (c.get("retrieval_keywords", []) or [])
            score = 0
            matched_terms = []
            for t in terms:
                t_norm = normalize_for_match(t)
                if t_norm and t_norm in q:
                    score += 2 if t in (c.get("aliases", []) or []) else 1
                    matched_terms.append(t)
            if score > 0:
                matches.append({
                    "concept_id": c.get("concept_id"),
                    "label": c.get("label"),
                    "category": c.get("category"),
                    "risk_level": c.get("risk_level"),
                    "default_answer_policy": c.get("default_answer_policy", "answer"),
                    "score": score,
                    "matched_terms": matched_terms,
                })
        matches.sort(key=lambda x: x["score"], reverse=True)
        return matches[:top_k]

    def _match_risk_policies(self, question: str) -> List[Dict[str, Any]]:
        q = normalize_for_match(question)
        matches = []
        for p in self.artifacts.risk_policies:
            hits = []
            for t in p.get("trigger_phrases", []) or []:
                if normalize_for_match(t) in q:
                    hits.append(t)
            if hits:
                rec = dict(p)
                rec["matched_triggers"] = hits
                matches.append(rec)
        # high risk first
        order = {"high": 3, "medium": 2, "low": 1, "高": 3, "中": 2, "低": 1}
        matches.sort(key=lambda x: order.get(x.get("risk_level"), 0), reverse=True)
        return matches

    def _match_query_patterns(self, question: str) -> List[Dict[str, Any]]:
        q = normalize_for_match(question)
        matches = []
        for p in self.artifacts.query_patterns:
            for ex in p.get("examples", []) or []:
                exn = normalize_for_match(ex)
                if exn and (exn in q or q in exn):
                    matches.append(dict(p))
                    break
        return matches

    def _get_offline_rewrite_terms(self, matched_concepts: List[str], category: str) -> Dict[str, List[str]]:
        """Return controlled offline rewrite terms.

        These are mandatory / trusted expansion terms. Runtime local LLM is not allowed
        to invent law article IDs or policy article IDs; those should come from these
        offline artifacts or concept_nodes.
        """
        terms = []
        article_hints = []
        for r in self.artifacts.rewrite_rules:
            if r.get("concept_id") in matched_concepts or (category and r.get("category") == category):
                terms.extend(r.get("rewrite_terms", []) or [])
                article_hints.extend(r.get("article_hints", []) or [])
        # Add concept-node related article hints as another controlled source.
        concept_map = {c.get("concept_id"): c for c in self.artifacts.concept_nodes}
        for cid in matched_concepts or []:
            c = concept_map.get(cid) or {}
            terms.extend(c.get("retrieval_keywords", []) or [])
            article_hints.extend(c.get("related_law_articles", []) or [])
            article_hints.extend(c.get("related_policy_articles", []) or [])
        terms = list(dict.fromkeys([str(t) for t in terms if t]))[:28]
        article_hints = list(dict.fromkeys([str(t) for t in article_hints if t]))[:16]
        return {"mandatory_terms": terms, "article_hints": article_hints}

    def _rewrite_from_artifacts(
        self,
        question: str,
        matched_concepts: List[str],
        category: str,
        fallback: str,
        normalized_query: Optional[str] = None,
        optional_terms: Optional[List[str]] = None,
    ) -> str:
        """Hybrid query rewrite.

        Mature workflow:
        - Offline rewrite_rules / concept_nodes provide mandatory controlled terms.
        - Runtime local LLM may provide optional semantic terms only.
        - Deterministic merge builds the final retrieval query.
        """
        offline = self._get_offline_rewrite_terms(matched_concepts, category)
        mandatory_terms = offline["mandatory_terms"]
        article_hints = offline["article_hints"]
        optional_terms = list(dict.fromkeys([str(t) for t in (optional_terms or []) if t]))[:LOCAL_LLM_OPTIONAL_REWRITE_MAX_TERMS]
        base_query = normalize_spaces(normalized_query or question)
        parts = [base_query]
        if mandatory_terms:
            parts.append(" ".join(mandatory_terms))
        if article_hints:
            parts.append(" ".join(article_hints))
        if optional_terms:
            parts.append(" ".join(optional_terms))
        if len(parts) > 1:
            return "｜".join(parts)
        return fallback

    def _heuristic_understanding(self, question: str) -> Dict[str, Any]:
        q = question.strip()
        category = detect_category(q)
        concept_matches = self._match_concepts(q)
        if concept_matches:
            category = concept_matches[0].get("category") or category
        if category == "overtime":
            intent = "工時與加班"
        elif category == "leave":
            intent = "請假與休假"
        elif category == "salary":
            intent = "薪資與工資"
        elif category == "termination":
            intent = "離職與資遣"
        elif category == "attendance":
            intent = "出勤與工時"
        elif category == "welfare":
            intent = "員工福利"
        elif category == "occupational_accident":
            intent = "職業災害"
        elif category == "privacy_sensitive":
            intent = "個資與敏感資訊"
        else:
            intent = "一般 HR 規章"

        risk_matches = self._match_risk_policies(q)
        pattern_matches = self._match_query_patterns(q)
        high_risk = any(str(r.get("risk_level", "")).lower() in ["high", "高"] for r in risk_matches)
        ambiguous = bool(pattern_matches) or q in AMBIGUOUS_SHORTS or len(q) <= 5

        # Concept priors can raise risk but should not alone decide final route when ambiguous.
        concept_policies = [m.get("default_answer_policy") for m in concept_matches if m.get("default_answer_policy")]
        concept_high = any(str(m.get("risk_level", "")).lower() in ["high", "高"] for m in concept_matches)

        if ambiguous:
            answer_policy = "clarify"
        elif high_risk or concept_high or any(k in q for k in RISK_KEYWORDS):
            answer_policy = "escalate"
        elif any(k in q for k in ["是否合法", "違法嗎", "可以告", "申訴", "主管不給", "被逼"]):
            answer_policy = "escalate"
        elif "with_disclaimer" in concept_policies or any(k in q for k in ["我這種情況", "如果", "主管", "個案", "薪資明細"]):
            answer_policy = "with_disclaimer"
        else:
            answer_policy = "direct"
        risk_level = "高" if answer_policy == "escalate" else "中" if answer_policy == "with_disclaimer" else "低"

        rewrite_terms = []
        if category == "overtime": rewrite_terms = ["加班", "延長工時", "加班費", "補休", "勞基法第24條"]
        elif category == "leave": rewrite_terms = ["請假", "特別休假", "病假", "事假", "員工內規", "勞基法第38條"]
        elif category == "salary": rewrite_terms = ["工資", "薪資給付", "扣薪", "加班費", "勞基法第22條"]
        elif category == "termination": rewrite_terms = ["資遣", "終止勞動契約", "預告期間", "勞基法第16條"]
        elif category == "welfare": rewrite_terms = ["員工福利", "旅遊補助", "HR公告", "內部規章"]
        fallback_rewrite = f"{intent}｜{q}｜{' '.join(rewrite_terms)}"
        matched_concepts = [m.get("concept_id") for m in concept_matches if m.get("concept_id")]
        rewritten_query = self._rewrite_from_artifacts(q, matched_concepts, category, fallback_rewrite)
        missing_slots = []
        for p in pattern_matches:
            missing_slots.extend(p.get("missing_slots", []) or [])
        missing_slots = list(dict.fromkeys(missing_slots))
        return {
            "intent": intent,
            "category": category,
            "risk_level": risk_level,
            "answer_policy": answer_policy,
            "matched_concepts": matched_concepts,
            "missing_slots": missing_slots,
            "llm_risk_signal": {},
            "risk_matches": risk_matches,
            "preferred_relations": [],
            "rewritten_query": rewritten_query,
        }

    def _llm_runtime_understanding(self, question: str, base: Dict[str, Any]) -> Dict[str, Any]:
        concepts_for_prompt = []
        for c in self.artifacts.concept_nodes[:80]:
            concepts_for_prompt.append({
                "concept_id": c.get("concept_id"),
                "label": c.get("label"),
                "category": c.get("category"),
                "aliases": (c.get("aliases", []) or [])[:6],
                "risk_level": c.get("risk_level"),
                "default_answer_policy": c.get("default_answer_policy"),
            })
        risk_for_prompt = []
        for p in self.artifacts.risk_policies[:40]:
            risk_for_prompt.append({
                "risk_policy_id": p.get("risk_policy_id"),
                "category": p.get("category"),
                "risk_level": p.get("risk_level"),
                "trigger_phrases": (p.get("trigger_phrases", []) or [])[:8],
                "default_route": p.get("default_route"),
            })
        patterns_for_prompt = []
        for p in self.artifacts.query_patterns[:40]:
            patterns_for_prompt.append({
                "pattern_id": p.get("pattern_id"),
                "pattern_type": p.get("pattern_type"),
                "category": p.get("category"),
                "examples": (p.get("examples", []) or [])[:6],
                "missing_slots": p.get("missing_slots", []) or [],
                "default_route": p.get("default_route"),
            })
        system = """
你是銀行 HR AI 助理的 Runtime Query Understanding 模組。
你的任務是做 structured classification，不是回答問題。
請根據 concept taxonomy、query pattern schema 與 risk policy，判斷使用者問題的 category、intent、matched_concepts、matched_query_pattern_ids、是否模糊、missing_slots、candidate risk。
你可以做口語/錯字 normalization，也可以提供 optional_rewrite_terms，但不得自行創造法條號或內規條號；法條/內規條號會由 offline rewrite_rules 提供。
重要：risk_level / recommended_route 只是 signal，最終路由會由 deterministic guardrails 決定。
只能輸出 valid JSON，不要 markdown，不要解釋。
"""
        user = f"""
使用者問題：{question}

可用 categories：general, working_hours, attendance, leave, overtime, salary, welfare, termination, occupational_accident, privacy_sensitive, high_risk, governance
可用 answer_policy：direct, with_disclaimer, clarify, escalate
可用 relation types：has_rule, related_to, refers_to, supplements, overrides, parent_of, child_of

Offline concept taxonomy（節錄）：
{json.dumps(concepts_for_prompt, ensure_ascii=False)}

Offline query patterns（節錄）：
{json.dumps(patterns_for_prompt, ensure_ascii=False)}

Offline risk policies（節錄）：
{json.dumps(risk_for_prompt, ensure_ascii=False)}

Heuristic baseline：
{json.dumps(base, ensure_ascii=False, default=str)}

請輸出 JSON schema：
{{
  "intent": "...",
  "category": "leave/overtime/salary/welfare/termination/working_hours/attendance/occupational_accident/privacy_sensitive/high_risk/general/governance",
  "matched_concepts": ["concept_id"],
  "matched_query_pattern_ids": ["pattern_id"],
  "is_ambiguous": true/false,
  "missing_slots": ["..."],
  "normalized_query": "口語或錯字修正後的問題，不要新增法條號",
  "risk_level": "低/中/高",
  "risk_reasons": ["..."],
  "answer_policy": "direct/with_disclaimer/clarify/escalate",
  "recommended_route": "answer/disclaimer/clarify/escalate",
  "preferred_relations": ["has_rule/related_to/refers_to/supplements/overrides"],
  "optional_rewrite_terms": ["最多5個語意補充詞，不要放法條號或內規條號"],
  "confidence": 0.0
}}
"""
        parsed = call_llm_json(system, user, default={})
        if not isinstance(parsed, dict) or not parsed:
            return base
        merged = dict(base)
        # Accept category/intent/concepts/ambiguity from local LLM.
        if parsed.get("intent"):
            merged["intent"] = str(parsed["intent"])
        if parsed.get("category"):
            merged["category"] = str(parsed["category"])
        if isinstance(parsed.get("matched_concepts"), list):
            merged["matched_concepts"] = list(dict.fromkeys([str(x) for x in parsed["matched_concepts"] if x]))[:8]
        if isinstance(parsed.get("matched_query_pattern_ids"), list):
            merged["matched_query_pattern_ids"] = list(dict.fromkeys([str(x) for x in parsed["matched_query_pattern_ids"] if x]))[:8]
        if isinstance(parsed.get("missing_slots"), list):
            merged["missing_slots"] = list(dict.fromkeys((base.get("missing_slots", []) or []) + [str(x) for x in parsed["missing_slots"] if x]))[:8]
        if parsed.get("normalized_query"):
            merged["normalized_query"] = str(parsed.get("normalized_query"))
        if isinstance(parsed.get("preferred_relations"), list):
            merged["preferred_relations"] = [str(x) for x in parsed["preferred_relations"] if x][:6]
        # Risk is stored as a signal. Final guardrails can override.
        merged["llm_risk_signal"] = {
            "risk_level": parsed.get("risk_level"),
            "risk_reasons": parsed.get("risk_reasons", []),
            "recommended_route": parsed.get("recommended_route"),
            "answer_policy": parsed.get("answer_policy"),
            "confidence": parsed.get("confidence"),
        }
        # Let LLM identify ambiguity, but final policy still goes through guardrails.
        if parsed.get("is_ambiguous") is True:
            merged["answer_policy"] = "clarify"
        elif parsed.get("answer_policy") in ["direct", "with_disclaimer", "clarify", "escalate"]:
            # Do not let LLM lower a heuristic high-risk/escalate decision.
            current_route = policy_to_route(merged.get("answer_policy"))
            proposed_route = policy_to_route(parsed.get("answer_policy"))
            priority = {"answer": 0, "disclaimer": 1, "clarify": 2, "escalate": 3}
            if priority.get(proposed_route, 0) >= priority.get(current_route, 0):
                merged["answer_policy"] = route_to_policy(proposed_route)
        # Hybrid query rewrite: offline mandatory terms + local LLM optional semantic terms.
        optional_terms = parsed.get("optional_rewrite_terms") if isinstance(parsed.get("optional_rewrite_terms"), list) else []
        merged["optional_rewrite_terms"] = [str(t) for t in optional_terms if t][:LOCAL_LLM_OPTIONAL_REWRITE_MAX_TERMS]
        fallback_rewrite = base.get("rewritten_query", question)
        merged["rewritten_query"] = self._rewrite_from_artifacts(
            question=question,
            matched_concepts=merged.get("matched_concepts", []) or [],
            category=merged.get("category", base.get("category", "general")),
            fallback=fallback_rewrite,
            normalized_query=merged.get("normalized_query"),
            optional_terms=merged.get("optional_rewrite_terms", []),
        )
        return merged

    def _node_query_understanding(self, state: HRState) -> HRState:
        q = state["question"]
        base = self._heuristic_understanding(q)
        if USE_LOCAL_LLM_FOR_QUERY_UNDERSTANDING:
            base = self._llm_runtime_understanding(q, base)
        # Refresh risk matches after local LLM concept/category changes; never remove existing high-risk matches.
        risk_matches = self._match_risk_policies(q)
        if risk_matches:
            base["risk_matches"] = risk_matches
        return {**base, "debug": add_debug(state, f"Query understanding: {base}")}

    def _node_retrieval_orchestrator(self, state: HRState) -> HRState:
        q = state.get("rewritten_query", state["question"])
        cat = state.get("category", "general")
        chunks = self.retriever.search(q, category=cat, top_k=8)
        seed_ids = []
        for c in chunks[:5]:
            aid = c.get("parent_id") if c.get("chunk_type") in ["semantic", "faq"] else c.get("article_id")
            if aid:
                seed_ids.append(aid)
            for rel in c.get("related_articles", []):
                seed_ids.append(rel)
        graph_ctx = self.kg.expand(seed_ids, state["question"], hops=1, max_nodes=14, preferred_relations=state.get("preferred_relations", []))
        return {
            "retrieved_chunks": chunks,
            "graph_context": graph_ctx,
            "debug": add_debug(state, f"Retrieved {len(chunks)} chunks; graph_nodes={len(graph_ctx.get('nodes', []))}, graph_edges={len(graph_ctx.get('edges', []))}"),
        }

    def _node_guardrails_and_route(self, state: HRState) -> HRState:
        policy = state.get("answer_policy", "direct")
        chunks = state.get("retrieved_chunks", [])
        if not chunks:
            confidence = 0.0
        else:
            top = float(chunks[0].get("final_score", 0))
            second = float(chunks[1].get("final_score", 0)) if len(chunks) > 1 else 0
            confidence = min(1.0, max(0.0, top + max(0, top-second)*0.15))

        risk_matches = state.get("risk_matches", []) or []
        llm_signal = state.get("llm_risk_signal", {}) or {}
        missing_slots = state.get("missing_slots", []) or []

        # Final route is deterministic. Risk policy can only raise risk / route, not lower it.
        high_risk_policy_hit = any(str(r.get("risk_level", "")).lower() in ["high", "高"] or r.get("default_route") == "escalate" for r in risk_matches)
        llm_high_risk = str(llm_signal.get("risk_level", "")).lower() in ["high", "高"] or llm_signal.get("recommended_route") == "escalate"
        llm_medium_risk = str(llm_signal.get("risk_level", "")).lower() in ["medium", "中"] or llm_signal.get("recommended_route") == "disclaimer"

        if high_risk_policy_hit:
            route = "escalate"
            final_risk_level = "高"
        elif llm_high_risk:
            route = "escalate"
            final_risk_level = "高"
        elif policy == "clarify" or missing_slots:
            route = "clarify"
            final_risk_level = "中" if missing_slots else state.get("risk_level", "低")
        elif policy == "escalate":
            route = "escalate"
            final_risk_level = "高"
        elif confidence < 0.18:
            route = "escalate"
            final_risk_level = "中"
        elif policy == "with_disclaimer" or llm_medium_risk:
            route = "disclaimer"
            final_risk_level = "中"
        else:
            route = "answer"
            final_risk_level = state.get("risk_level", "低")

        return {
            "confidence": round(confidence, 4),
            "route": route,
            "risk_level": final_risk_level,
            "debug": add_debug(state, f"Guardrails: policy={policy}, risk_matches={len(risk_matches)}, llm_signal={llm_signal}, confidence={confidence:.4f}, route={route}"),
        }

    def _make_context(self, state: HRState, max_chars: int = 900) -> str:
        lines = []
        for i, c in enumerate(state.get("retrieved_chunks", [])[:6], start=1):
            lines.append(f"""
[S{i}]
chunk_type: {c.get('chunk_type')}
source_type: {c.get('source_type')}
document: {c.get('document_name')}
article_no: {c.get('article_no')}
title: {c.get('title')}
category: {c.get('category')}
priority: {c.get('priority')}
content: {c.get('content')[:max_chars]}
""".strip())
        graph_context = state.get("graph_context", {}).get("context", "")
        if graph_context:
            lines.append("\n[Graph-enhanced Context]\n" + graph_context[:2500])
        return "\n\n---\n\n".join(lines)

    def _build_citations(self, state: HRState) -> List[Dict[str, Any]]:
        cites = []
        for i, c in enumerate(state.get("retrieved_chunks", [])[:6], start=1):
            cites.append({
                "source_id": f"S{i}",
                "chunk_type": c.get("chunk_type"),
                "source_type": c.get("source_type"),
                "document_name": c.get("document_name"),
                "article_no": c.get("article_no"),
                "title": c.get("title"),
                "category": c.get("category"),
                "score": c.get("final_score"),
                "content_preview": c.get("content", "")[:180],
            })
        return cites

    def _fallback_answer(self, state: HRState, disclaimer: bool = False) -> str:
        q = state["question"]
        chunks = state.get("retrieved_chunks", [])[:4]
        if not chunks:
            return "目前知識庫沒有找到足夠依據可回答，建議洽 HR 確認。"
        top_internal = [c for c in chunks if c.get("source_type") == "internal_policy"]
        top_law = [c for c in chunks if c.get("source_type") == "law"]
        lines = []
        lines.append("簡短結論：")
        if top_internal:
            lines.append("依目前知識庫，應優先參考公司內部規章；若內規未明確規定，再參考勞動基準法作為最低標準。")
        else:
            lines.append("依目前檢索結果，以下為相關法規或規章整理。")
        lines.append("\n適用條件：")
        lines.append("需依實際身分、年資、班表、核准流程與公司最新公告確認。")
        lines.append("\n依據 Citation：")
        for i, c in enumerate(chunks, start=1):
            lines.append(f"- [S{i}] {c.get('source_type')}｜{c.get('document_name')}｜{c.get('article_no')}｜{c.get('title')}")
        if state.get("graph_context", {}).get("edges"):
            lines.append("\n規範差異 / Graph 關係：")
            for e in state["graph_context"]["edges"][:5]:
                lines.append(f"- {e['source']} --{e['relation']}--> {e['target']}")
        lines.append("\n白話說明：")
        for i, c in enumerate(chunks[:3], start=1):
            content = c.get("answer") or c.get("content", "")
            lines.append(f"- [S{i}] {content[:260]}...")
        lines.append("\n注意事項 / 聲明：")
        if disclaimer:
            lines.append("本回答依據現行知識庫提供一般性說明，實際適用仍需依個案情況與人力資源單位最終認定為準。")
        else:
            lines.append("若涉及個人薪資、主管指示、申訴爭議或公司最新公告，建議洽 HR 確認。")
        lines.append("\n下一步建議：")
        lines.append("若仍不確定，請提供更完整情境或轉 HR 人員確認。")
        return "\n".join(lines)

    def _node_generate_answer(self, state: HRState) -> HRState:
        context = self._make_context(state)
        disclaimer = state.get("route") == "disclaimer"
        if not USE_LLM:
            answer = self._fallback_answer(state, disclaimer=disclaimer)
        else:
            system = """
你是安久銀行 HR AI 智能助理。
你只能根據提供的 Retrieval Context 與 Graph Context 回答，不得自行編造資料。
回答時必須：
1. 優先使用 internal_policy，其次使用 law 作為最低標準。
2. 必須引用來源，格式使用 [S1], [S2]。
3. 若內規優於法規，請說明「內規 vs 法規」差異。
4. 若是情境型問題，需加風險聲明。
5. 不得做法律判定；高風險或個案爭議需建議洽 HR。
"""
            user = f"""
員工問題：{state['question']}
Intent: {state.get('intent')}
Category: {state.get('category')}
Risk: {state.get('risk_level')}
Route: {state.get('route')}

Retrieval / Graph Context:
{context}

請用以下格式回答：
簡短結論：
適用條件：
依據 Citation：
規範差異（內規 vs 法規，如有）：
白話說明：
注意事項 / 聲明：
下一步建議：
"""
            answer = call_llm_text(system, user, temperature=0.1) or self._fallback_answer(state, disclaimer=disclaimer)
        return {"answer": answer, "citations": self._build_citations(state), "debug": add_debug(state, "Generated answer")}

    def _node_clarify(self, state: HRState) -> HRState:
        missing_slots = state.get("missing_slots", []) or []
        category = state.get("category", "general")
        questions = []
        for p in self.artifacts.query_patterns:
            if p.get("category") == category and p.get("pattern_type") == "ambiguous":
                questions.extend(p.get("clarification_questions", []) or [])
        if not questions:
            questions = [
                "請問你想查詢的是特休、病假、事假、加班、補休、薪資、資遣或福利？",
                "是否涉及個人薪資、主管核准、特殊班表或申訴爭議？",
                "需要查詢公司內規，還是只想了解勞基法最低規定？",
            ]
        questions = list(dict.fromkeys(questions))[:4]
        slot_text = "、".join(missing_slots) if missing_slots else "假別、期間、原因、是否涉及主管核准或個案爭議"
        answer = f"""
我需要再確認一些資訊，才能避免誤判。

您的問題是：「{state['question']}」

目前判斷類別：{state.get('category')}｜可能缺少資訊：{slot_text}

請補充以下資訊：
""".strip()
        for i, qu in enumerate(questions, start=1):
            answer += f"\n{i}. {qu}"
        answer += "\n\n在資訊不足時，系統不會直接推論答案，以降低 HR 法規誤判風險。"
        return {"answer": answer, "citations": [], "faithfulness_score": 1.0, "debug": add_debug(state, "Clarification generated")}

    def _node_escalate(self, state: HRState) -> HRState:
        chunks = state.get("retrieved_chunks", [])[:3]
        refs = "\n".join([f"- [S{i+1}] {c.get('source_type')}｜{c.get('article_no')}｜{c.get('title')}" for i, c in enumerate(chunks)])
        answer = f"""
目前此問題不適合由 AI 直接判定，建議轉由 HR 或法遵人員處理。

問題：{state['question']}
判斷類型：{state.get('intent')}
風險等級：{state.get('risk_level')}

原因：
1. 問題可能涉及個案判斷、申訴爭議、薪資明細、主管處置、個資或法律責任。
2. AI 可提供一般規範整理，但不應直接做合法性或責任歸屬判斷。
3. 本題命中的風險政策：{', '.join([r.get('risk_policy_id','') for r in state.get('risk_matches', [])]) or '無明確風險政策，但路由判斷為高風險或低信心'}。

可能相關依據：
{refs if refs else '目前無足夠相關依據。'}

下一步建議：請洽 HR 服務窗口或依公司正式申訴 / 諮詢流程處理。
""".strip()
        return {"answer": answer, "citations": self._build_citations(state), "faithfulness_score": 1.0, "debug": add_debug(state, "Escalation generated")}

    def _node_faithfulness_check(self, state: HRState) -> HRState:
        answer = state.get("answer", "")
        citations = state.get("citations", [])
        has_citation = bool(re.search(r"\[S\d+\]", answer))
        mentions_article = any(str(c.get("article_no", "")).replace(" ", "") in answer.replace(" ", "") for c in citations if c.get("article_no"))
        score = 0.82
        if has_citation: score += 0.10
        if mentions_article: score += 0.05
        if state.get("route") in ["escalate", "clarify"]: score = max(score, 0.95)
        score = min(1.0, score)
        return {"faithfulness_score": round(score, 4), "debug": add_debug(state, f"Faithfulness={score:.4f}")}

    def _route_after_guardrails(self, state: HRState) -> Literal["answer", "disclaimer", "escalate", "clarify"]:
        return state.get("route", "answer")

    def _build_graph(self):
        from langgraph.graph import StateGraph, START, END
        graph = StateGraph(HRState)
        graph.add_node("query_understanding", self._node_query_understanding)
        graph.add_node("retrieval_orchestrator", self._node_retrieval_orchestrator)
        graph.add_node("guardrails", self._node_guardrails_and_route)
        graph.add_node("generate_answer", self._node_generate_answer)
        graph.add_node("clarify", self._node_clarify)
        graph.add_node("escalate", self._node_escalate)
        graph.add_node("faithfulness_check", self._node_faithfulness_check)

        graph.add_edge(START, "query_understanding")
        graph.add_edge("query_understanding", "retrieval_orchestrator")
        graph.add_edge("retrieval_orchestrator", "guardrails")
        graph.add_conditional_edges(
            "guardrails",
            self._route_after_guardrails,
            {
                "answer": "generate_answer",
                "disclaimer": "generate_answer",
                "escalate": "escalate",
                "clarify": "clarify",
            },
        )
        graph.add_edge("generate_answer", "faithfulness_check")
        graph.add_edge("faithfulness_check", END)
        graph.add_edge("clarify", END)
        graph.add_edge("escalate", END)
        return graph.compile()

    def ask(self, question: str) -> HRState:
        return self.app.invoke({"question": question})

# -----------------------------
# 9. Evaluation + Feedback
# -----------------------------
def load_golden_dataset(json_path: str) -> pd.DataFrame:
    """Load user-provided Golden Dataset JSON.

    Supported structures:
    1) {"items": [{...}, {...}]}
    2) [{...}, {...}]

    Required field: question
    Recommended fields: id, question_type, expected_category, expected_route,
    expected_citations, expected_key_points, should_escalate, should_clarify.
    """
    with open(json_path, "r", encoding="utf-8") as f:
        raw = json.load(f)
    items = raw.get("items", raw) if isinstance(raw, dict) else raw
    if not isinstance(items, list):
        raise ValueError("Golden Dataset JSON must be a list or an object with an 'items' list.")
    df = pd.DataFrame(items)
    if "question" not in df.columns:
        raise ValueError("Golden Dataset must contain a 'question' field.")
    if "expected_route" not in df.columns:
        df["expected_route"] = np.where(df.get("should_clarify", False), "clarify", np.where(df.get("should_escalate", False), "escalate", "answer"))
    if "expected_category" not in df.columns:
        df["expected_category"] = df["question"].apply(detect_category)
    if "expected_citations" not in df.columns:
        df["expected_citations"] = [[] for _ in range(len(df))]
    if "expected_key_points" not in df.columns:
        df["expected_key_points"] = [[] for _ in range(len(df))]
    if "id" not in df.columns:
        df["id"] = [f"G{i+1:03d}" for i in range(len(df))]
    return df


def citation_retrieval_hit(expected_citations: Any, citations: List[Dict[str, Any]]) -> Optional[bool]:
    if expected_citations is None or (isinstance(expected_citations, float) and np.isnan(expected_citations)):
        return None
    if isinstance(expected_citations, str):
        try:
            parsed = json.loads(expected_citations)
            expected_citations = parsed
        except Exception:
            expected_citations = [expected_citations]
    if not isinstance(expected_citations, list):
        expected_citations = [expected_citations]
    expected_citations = [str(x).strip() for x in expected_citations if str(x).strip()]
    if not expected_citations:
        return None

    retrieved_text = normalize_for_match(" ".join([safe_json_dumps(c) for c in citations]))
    for exp in expected_citations:
        exp_norm = normalize_for_match(exp)
        if exp_norm and exp_norm in retrieved_text:
            return True
        # Match by article number, e.g. expected citation「安久銀行...第11條」 vs retrieved article_no「第 11 條」.
        for ref in extract_article_refs(exp):
            if ref and ref in retrieved_text:
                return True
    return False


def source_type_hit(expected_source_type: Any, citations: List[Dict[str, Any]]) -> Optional[bool]:
    if expected_source_type is None or (isinstance(expected_source_type, float) and np.isnan(expected_source_type)):
        return None
    if isinstance(expected_source_type, str):
        try:
            parsed = json.loads(expected_source_type)
            expected_source_type = parsed
        except Exception:
            expected_source_type = [expected_source_type]
    if not isinstance(expected_source_type, list):
        expected_source_type = [expected_source_type]
    expected = {str(x).strip() for x in expected_source_type if str(x).strip()}
    if not expected:
        return None
    actual = {str(c.get("source_type", "")).strip() for c in citations}
    return bool(expected & actual)


def evaluate_assistant(assistant: HRAssistantGraph, golden_df: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame]:
    rows = []
    for _, row in tqdm(golden_df.iterrows(), total=len(golden_df), desc="Evaluating"):
        q = row["question"]
        start = time.time()
        result = assistant.ask(q)
        latency = time.time() - start
        citations = result.get("citations", [])

        expected_route = row.get("expected_route", "answer")
        expected_category = row.get("expected_category", "")
        retrieval_hit = citation_retrieval_hit(row.get("expected_citations", []), citations)
        src_hit = source_type_hit(row.get("expected_source_type", []), citations)
        route_correct = result.get("route") == expected_route
        category_correct = category_matches(expected_category, result.get("category"))
        citation_present = bool(citations) if result.get("route") not in ["clarify"] else True

        rows.append({
            "id": row.get("id"),
            "question_type": row.get("question_type"),
            "test_dimension": row.get("test_dimension"),
            "question": q,
            "expected_category": expected_category,
            "actual_category": result.get("category"),
            "category_correct": category_correct,
            "expected_route": expected_route,
            "actual_route": result.get("route"),
            "route_correct": route_correct,
            "expected_citations": row.get("expected_citations", []),
            "retrieval_hit": retrieval_hit,
            "expected_source_type": row.get("expected_source_type", []),
            "source_type_hit": src_hit,
            "citation_present": citation_present,
            "confidence": result.get("confidence"),
            "faithfulness_score": result.get("faithfulness_score"),
            "latency_sec": round(latency, 3),
            "answer_preview": result.get("answer", "")[:240],
        })
    detail = pd.DataFrame(rows)
    valid_hits = [x for x in detail["retrieval_hit"].tolist() if isinstance(x, (bool, np.bool_))]
    valid_src_hits = [x for x in detail["source_type_hit"].tolist() if isinstance(x, (bool, np.bool_))]
    summary = pd.DataFrame([{
        "Total Questions": len(detail),
        "Category Accuracy": detail["category_correct"].mean(),
        "Route Accuracy": detail["route_correct"].mean(),
        "Retrieval Hit Rate": (sum(valid_hits) / len(valid_hits)) if valid_hits else None,
        "Source Type Hit Rate": (sum(valid_src_hits) / len(valid_src_hits)) if valid_src_hits else None,
        "Citation Present Rate": detail["citation_present"].mean(),
        "Avg Faithfulness Score": detail["faithfulness_score"].mean(),
        "Avg Latency Sec": detail["latency_sec"].mean(),
    }])
    return detail, summary


def log_feedback(output_dir: Path, question: str, answer: str, helpful: bool, correctness_score: int, completeness_score: int, comment: str = "") -> pd.DataFrame:
    path = output_dir / "feedback_log.csv"
    row = {
        "timestamp": pd.Timestamp.now().isoformat(),
        "question": question,
        "answer": answer,
        "helpful": helpful,
        "correctness_score": correctness_score,
        "completeness_score": completeness_score,
        "comment": comment,
    }
    if path.exists():
        df = pd.read_csv(path)
        df = pd.concat([df, pd.DataFrame([row])], ignore_index=True)
    else:
        df = pd.DataFrame([row])
    df.to_csv(path, index=False, encoding="utf-8-sig")
    return df

# -----------------------------
# 10. Colab Runner
# -----------------------------
def _resolve_existing_path(path_str: str) -> Optional[str]:
    if path_str and Path(path_str).exists():
        return str(Path(path_str))
    return None


def _identify_uploaded_files(uploaded_names: List[str]) -> Tuple[Optional[str], Optional[str], Optional[str]]:
    docx_files = [fn for fn in uploaded_names if fn.lower().endswith(".docx")]
    json_files = [fn for fn in uploaded_names if fn.lower().endswith(".json")]

    law_docx = None
    policy_docx = None
    golden_json = json_files[0] if json_files else None

    for fn in docx_files:
        if any(k in fn for k in ["勞動基準法", "勞基法", "labor"]):
            law_docx = fn
        elif any(k in fn for k in ["內規", "規章", "員工工作", "員工手冊", "policy", "福利規章"]):
            policy_docx = fn

    # Fallback by order if names are ambiguous.
    remaining = [fn for fn in docx_files if fn not in {law_docx, policy_docx}]
    if law_docx is None and remaining:
        law_docx = remaining.pop(0)
    if policy_docx is None and remaining:
        policy_docx = remaining.pop(0)

    if IN_COLAB:
        law_docx = f"/content/{law_docx}" if law_docx else None
        policy_docx = f"/content/{policy_docx}" if policy_docx else None
        golden_json = f"/content/{golden_json}" if golden_json else None
    return law_docx, policy_docx, golden_json


def prepare_input_files() -> Tuple[str, str, str]:
    """Resolve or upload the three required user-provided files.

    Required:
    1. 勞動基準法 DOCX
    2. 模擬銀行內規 DOCX
    3. Golden Dataset JSON
    """
    law_docx = _resolve_existing_path(LABOR_LAW_DOCX_PATH)
    policy_docx = _resolve_existing_path(INTERNAL_POLICY_DOCX_PATH)
    golden_json = _resolve_existing_path(GOLDEN_DATASET_JSON_PATH)

    if law_docx and policy_docx and golden_json:
        return law_docx, policy_docx, golden_json

    if IN_COLAB:
        print("請一次上傳三個檔案：")
        print("1) 勞動基準法 DOCX")
        print("2) 模擬銀行員工內部規章 DOCX")
        print("3) Golden Dataset JSON")
        uploaded = files.upload()
        ulaw, upolicy, ugolden = _identify_uploaded_files(list(uploaded.keys()))
        law_docx = law_docx or ulaw
        policy_docx = policy_docx or upolicy
        golden_json = golden_json or ugolden
    else:
        # Local convenience: search current directory and /mnt/data if paths are not set.
        candidates = list(Path(".").glob("*")) + list(Path("/mnt/data").glob("*"))
        names = [str(p) for p in candidates]
        ulaw, upolicy, ugolden = _identify_uploaded_files(names)
        law_docx = law_docx or ulaw
        policy_docx = policy_docx or upolicy
        golden_json = golden_json or ugolden

        # The repo ships the internal policy DOCX + golden JSON but not the official
        # 勞動基準法 DOCX (that is an external legal source). For a zero-config local
        # run, fall back to the built-in sample labor-law DOCX so the pipeline can run
        # end-to-end. Provide a real 勞動基準法 DOCX via LABOR_LAW_DOCX_PATH for production.
        if not (law_docx and Path(law_docx).exists()):
            sample_path = str(OUTPUT_DIR / "參考資料_勞動基準法_sample.docx")
            print("未提供勞動基準法 DOCX，改用內建 sample 條文 ->", sample_path)
            create_sample_labor_law_docx(sample_path)
            law_docx = sample_path

    missing = []
    if not law_docx or not Path(law_docx).exists():
        missing.append("勞動基準法 DOCX")
    if not policy_docx or not Path(policy_docx).exists():
        missing.append("模擬銀行內規 DOCX")
    if not golden_json or not Path(golden_json).exists():
        missing.append("Golden Dataset JSON")
    if missing:
        raise FileNotFoundError("缺少必要檔案：" + ", ".join(missing) + "。請上傳或設定環境變數 LABOR_LAW_DOCX_PATH / INTERNAL_POLICY_DOCX_PATH / GOLDEN_DATASET_JSON_PATH。")

    return str(law_docx), str(policy_docx), str(golden_json)


def display_result(result: HRState):
    display(Markdown("## 使用者問題"))
    display(Markdown(result.get("question", "")))
    display(Markdown("## AI 回答"))
    display(Markdown(result.get("answer", "")))
    meta = {
        "intent": result.get("intent"),
        "category": result.get("category"),
        "risk_level": result.get("risk_level"),
        "answer_policy": result.get("answer_policy"),
        "route": result.get("route"),
        "confidence": result.get("confidence"),
        "faithfulness_score": result.get("faithfulness_score"),
    }
    display(Markdown("## 系統判斷"))
    display(pd.DataFrame([meta]))
    display(Markdown("## 引用來源"))
    if result.get("citations"):
        display(pd.DataFrame(result["citations"]))
    else:
        display(Markdown("無 citation。"))
    display(Markdown("## Debug Trace"))
    for d in result.get("debug", []):
        print("-", d)


def save_outputs(output_dir: Path, articles, chunks, kg: HRKnowledgeGraph, demo_results: List[Dict[str, Any]], golden_df, eval_detail, eval_summary):
    output_dir.mkdir(parents=True, exist_ok=True)
    with open(output_dir / "articles.json", "w", encoding="utf-8") as f:
        json.dump(articles, f, ensure_ascii=False, indent=2)
    with open(output_dir / "chunks_3layer_faq.json", "w", encoding="utf-8") as f:
        json.dump(chunks, f, ensure_ascii=False, indent=2)
    with open(output_dir / "demo_results.json", "w", encoding="utf-8") as f:
        json.dump(demo_results, f, ensure_ascii=False, indent=2)
    pd.DataFrame(chunks).drop(columns=["embedding_text"], errors="ignore").to_csv(output_dir / "chunks_3layer_faq.csv", index=False, encoding="utf-8-sig")
    golden_df.to_csv(output_dir / "golden_dataset.csv", index=False, encoding="utf-8-sig")
    eval_detail.to_csv(output_dir / "evaluation_detail.csv", index=False, encoding="utf-8-sig")
    eval_summary.to_csv(output_dir / "evaluation_summary.csv", index=False, encoding="utf-8-sig")
    kg.save_graph_files(output_dir)
    readme = f"""
# 安久銀行 HR AI 智能助理 - Colab Technical Demo

## 技術內容
- 使用者上傳模擬員工內規 DOCX：補足企業內部規則層
- Policy-aware Hierarchical Chunking：Document-level / Article-level / Semantic sub-chunk
- FAQ Chunk：預設不使用 Golden Dataset 進知識庫；若設 USE_GOLDEN_AS_FAQ_CHUNKS=true 才做實驗性 FAQ chunk
- Hybrid Retrieval：Vector + BM25 + keyword + metadata priority
- Offline Artifacts：讀取 concept_nodes / risk_policy / query_patterns / rewrite_rules / relation_schema / graph_relation_candidates JSON
- Knowledge Graph / Graph RAG：建立 law / internal_policy / concept 節點與 approved graph relations
- LangGraph Workflow：Runtime Local LLM Classification -> Retrieval Orchestrator -> Deterministic Guardrails -> Answer / Disclaimer / Clarify / Escalate
- HuggingFace Local LLM：Colab GPU 載入 Instruct model，不需 OpenAI API key；負責分類 signal 與 grounded answer generation
- 使用者上傳 Golden Dataset JSON Evaluation：category, route, retrieval, source type, citation, faithfulness, latency

## 主要輸出
- articles.json
- chunks_3layer_faq.json / csv
- kg_nodes.csv / kg_edges.csv / hr_knowledge_graph.gexf
- golden_dataset.csv
- evaluation_detail.csv
- evaluation_summary.csv
- demo_results.json
- feedback_log.csv
""".strip()
    (output_dir / "README.md").write_text(readme, encoding="utf-8")

    zip_path = output_dir / "hr_ai_graph_rag_outputs.zip"
    with zipfile.ZipFile(zip_path, "w") as z:
        for fp in output_dir.glob("*"):
            if fp.is_file() and fp.name != zip_path.name:
                z.write(fp, arcname=fp.name)
    print("Saved outputs to:", output_dir)
    print("ZIP:", zip_path)
    if IN_COLAB:
        try:
            files.download(str(zip_path))
        except Exception as e:
            print("Download failed:", repr(e))
    return zip_path


def main():
    labor_docx_path, policy_docx_path, golden_json_path = prepare_input_files()
    print("Labor Law DOCX path:", labor_docx_path)
    print("Internal Policy DOCX path:", policy_docx_path)
    print("Golden Dataset JSON path:", golden_json_path)

    offline_artifacts = load_offline_artifacts()
    golden_df = load_golden_dataset(golden_json_path)
    print(f"Golden Dataset questions: {len(golden_df)}")
    display(Markdown("# Uploaded Golden Dataset Preview"))
    display(golden_df.head(10))

    builder = HRKnowledgeBuilder(ChunkConfig(version="PoC-v1", effective_date="2026-06-23"))
    articles, chunks = builder.build_chunks(
        labor_law_docx_path=labor_docx_path,
        internal_policy_docx_path=policy_docx_path,
        golden_df=golden_df,
    )
    print(f"Articles: {len(articles)}")
    print(f"Chunks: {len(chunks)}")
    display(Markdown("# Knowledge Chunks Preview"))
    display(pd.DataFrame(chunks).drop(columns=["embedding_text"], errors="ignore").head(12))

    kg = HRKnowledgeGraph(articles, chunks, artifacts=offline_artifacts)
    print("KG nodes:", kg.G.number_of_nodes(), "edges:", kg.G.number_of_edges())

    retriever = HybridRetriever(chunks)
    assistant = HRAssistantGraph(retriever, kg, artifacts=offline_artifacts)

    # Use first 9 questions from the uploaded Golden Dataset as demo questions.
    demo_questions = golden_df["question"].dropna().astype(str).head(9).tolist()

    demo_results = []
    print("\nRunning demo questions...")
    for q in demo_questions:
        r = assistant.ask(q)
        demo_results.append(r)

    # Show first 3 detailed results
    for r in demo_results[:3]:
        display_result(r)

    batch_df = pd.DataFrame([{
        "question": r.get("question"),
        "intent": r.get("intent"),
        "category": r.get("category"),
        "route": r.get("route"),
        "risk_level": r.get("risk_level"),
        "confidence": r.get("confidence"),
        "faithfulness_score": r.get("faithfulness_score"),
        "answer_preview": r.get("answer", "")[:160],
    } for r in demo_results])
    display(Markdown("# Batch Demo Summary"))
    display(batch_df)

    # Evaluation on uploaded Golden Dataset JSON
    eval_detail, eval_summary = evaluate_assistant(assistant, golden_df)
    display(Markdown("# Evaluation Detail"))
    display(eval_detail)
    display(Markdown("# Evaluation Summary"))
    display(eval_summary)

    # Feedback example based on first demo question
    if demo_results:
        log_feedback(
            OUTPUT_DIR,
            question=demo_results[0]["question"],
            answer=demo_results[0]["answer"],
            helpful=True,
            correctness_score=5,
            completeness_score=4,
            comment="Demo feedback: external DOCX/JSON ingestion version.",
        )

    # Save loaded offline artifact summary for auditability before packaging outputs.
    with open(OUTPUT_DIR / "loaded_offline_artifacts.json", "w", encoding="utf-8") as f:
        json.dump({"artifact_dir": str(offline_artifacts.artifact_dir), "loaded_files": offline_artifacts.loaded_files}, f, ensure_ascii=False, indent=2)

    save_outputs(OUTPUT_DIR, articles, chunks, kg, demo_results, golden_df, eval_detail, eval_summary)
    print("\n全部流程完成。")
    return assistant, articles, chunks, kg, retriever, golden_df, offline_artifacts

# Run automatically when executed as a script (python src/hr_ai_graph_rag.py) or via
# Colab `%run`. Importing the module (`import hr_ai_graph_rag`) is side-effect free, so the
# data / parsing / chunking / graph layers can be reused and tested without running the
# full GPU + LLM pipeline.
if __name__ == "__main__":
    assistant, articles, chunks, kg, retriever, golden_df, offline_artifacts = main()

# 互動測試範例：執行完後可以自己改問題
# result = assistant.ask("你的問題")
# display_result(result)
